#!/usr/bin/env python3
"""
Sharedo Batch DOCX to HTML Converter
Processes multiple documents with confidence scoring and comprehensive reporting
"""

import os
import re
import json
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
from datetime import datetime
from docx import Document
from bs4 import BeautifulSoup
import traceback

class SharedoBatchConverter:
    """Batch converter with confidence scoring and reporting"""
    
    def __init__(self, input_folder="Input", output_folder="Output"):
        self.input_folder = Path(input_folder)
        self.output_folder = Path(output_folder)
        self.report_data = {
            "conversion_date": datetime.now().isoformat(),
            "total_files": 0,
            "successful": 0,
            "failed": 0,
            "files": []
        }
        
        # Ensure folders exist
        self.input_folder.mkdir(exist_ok=True)
        self.output_folder.mkdir(exist_ok=True)
        
        # Known Sharedo patterns for detection
        self.sharedo_patterns = {
            'content_control': re.compile(r'«([^»]+)»'),
            'placeholder': re.compile(r'\[_+\]'),
            'context_var': re.compile(r'context\.[a-zA-Z0-9._!?&=]+'),
            'document_var': re.compile(r'document\.[a-zA-Z0-9._!?&=]+'),
            'conditional_marker': re.compile(r'#if|#foreach|#else|#endif'),
            'sharedo_tag': re.compile(r'Sharedo\s+(Tag|Section|ContentBlock):'),
        }
    
    def process_all_documents(self):
        """Process all DOCX files in input folder"""
        print("=" * 70)
        print("🚀 SHAREDO BATCH CONVERTER")
        print("=" * 70)
        
        # Find all DOCX files
        docx_files = list(self.input_folder.glob("*.docx"))
        
        if not docx_files:
            print(f"⚠️ No DOCX files found in {self.input_folder}")
            return
        
        print(f"📁 Found {len(docx_files)} DOCX files to process")
        print("-" * 70)
        
        self.report_data["total_files"] = len(docx_files)
        
        # Process each file
        for docx_file in docx_files:
            # Skip temporary files
            if docx_file.name.startswith('~$'):
                continue
                
            print(f"\n📄 Processing: {docx_file.name}")
            file_report = self.process_single_document(docx_file)
            self.report_data["files"].append(file_report)
            
            if file_report["status"] == "success":
                self.report_data["successful"] += 1
                print(f"   ✅ Success - Confidence: {file_report['confidence_score']}%")
            else:
                self.report_data["failed"] += 1
                print(f"   ❌ Failed: {file_report.get('error', 'Unknown error')}")
        
        # Generate final report
        self.generate_final_report()
        
        print("\n" + "=" * 70)
        print("📊 CONVERSION SUMMARY")
        print("=" * 70)
        print(f"Total Files: {self.report_data['total_files']}")
        print(f"Successful: {self.report_data['successful']}")
        print(f"Failed: {self.report_data['failed']}")
        print(f"Report saved: {self.output_folder}/conversion_report.html")
        print("=" * 70)
    
    def process_single_document(self, docx_path):
        """Process a single DOCX file with confidence scoring"""
        file_report = {
            "filename": docx_path.name,
            "status": "pending",
            "confidence_score": 100,
            "issues": [],
            "warnings": [],
            "statistics": {},
            "sharedo_elements": {},
            "requires_review": False
        }
        
        try:
            # Analyze document first
            analysis = self.analyze_document(docx_path)
            file_report["statistics"] = analysis["statistics"]
            file_report["sharedo_elements"] = analysis["sharedo_elements"]
            
            # Calculate confidence score
            confidence, issues, warnings = self.calculate_confidence(analysis)
            file_report["confidence_score"] = confidence
            file_report["issues"] = issues
            file_report["warnings"] = warnings
            file_report["requires_review"] = confidence < 90 or len(issues) > 0
            
            # Convert document
            html_content = self.convert_document(docx_path, analysis)
            
            # Save HTML
            output_path = self.output_folder / f"{docx_path.stem}.html"
            output_path.write_text(html_content, encoding='utf-8')
            
            file_report["status"] = "success"
            file_report["output_file"] = str(output_path)
            
        except Exception as e:
            file_report["status"] = "failed"
            file_report["error"] = str(e)
            file_report["traceback"] = traceback.format_exc()
            file_report["requires_review"] = True
            file_report["confidence_score"] = 0
        
        return file_report
    
    def analyze_document(self, docx_path):
        """Analyze document for Sharedo elements and complexity"""
        analysis = {
            "statistics": {
                "paragraphs": 0,
                "tables": 0,
                "images": 0,
                "total_words": 0
            },
            "sharedo_elements": {
                "content_controls": [],
                "tags": [],
                "sections": [],
                "content_blocks": [],
                "conditionals": [],
                "placeholders": []
            },
            "complexity_indicators": {
                "has_nested_tables": False,
                "has_complex_formatting": False,
                "has_custom_styles": False,
                "has_macros": False
            }
        }
        
        # Extract content controls from XML
        with zipfile.ZipFile(docx_path, 'r') as docx_zip:
            if 'word/document.xml' in docx_zip.namelist():
                doc_xml = docx_zip.read('word/document.xml').decode('utf-8')
                
                # Parse XML for content controls
                namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                root = ET.fromstring(doc_xml)
                
                # Find all SDT elements (content controls)
                for sdt in root.findall('.//w:sdt', namespaces):
                    sdt_pr = sdt.find('w:sdtPr', namespaces)
                    if sdt_pr is not None:
                        tag_elem = sdt_pr.find('w:tag', namespaces)
                        alias_elem = sdt_pr.find('w:alias', namespaces)
                        
                        if tag_elem is not None:
                            tag = tag_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                            alias = alias_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') if alias_elem is not None else None
                            
                            control_info = {"tag": tag, "alias": alias}
                            analysis["sharedo_elements"]["content_controls"].append(control_info)
                            
                            # Categorize by type
                            if alias:
                                if 'ContentBlock' in alias:
                                    analysis["sharedo_elements"]["content_blocks"].append(tag)
                                elif 'Section' in alias:
                                    analysis["sharedo_elements"]["sections"].append(tag)
                                elif 'Tag' in alias:
                                    analysis["sharedo_elements"]["tags"].append(tag)
        
        # Analyze document content
        doc = Document(docx_path)
        
        # Statistics
        analysis["statistics"]["paragraphs"] = len(doc.paragraphs)
        analysis["statistics"]["tables"] = len(doc.tables)
        
        # Word count and pattern detection
        full_text = ""
        for para in doc.paragraphs:
            text = para.text
            full_text += text + " "
            
            # Check for patterns
            if self.sharedo_patterns['placeholder'].search(text):
                analysis["sharedo_elements"]["placeholders"].append(text[:50])
            
            if self.sharedo_patterns['conditional_marker'].search(text):
                analysis["sharedo_elements"]["conditionals"].append(text[:50])
            
            # Check for custom styles
            if para.style and para.style.name not in ['Normal', 'Heading 1', 'Heading 2', 'Heading 3']:
                analysis["complexity_indicators"]["has_custom_styles"] = True
        
        analysis["statistics"]["total_words"] = len(full_text.split())
        
        # Check for nested tables (complexity indicator)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if len(cell.tables) > 0:
                        analysis["complexity_indicators"]["has_nested_tables"] = True
                        break
        
        return analysis
    
    def calculate_confidence(self, analysis):
        """Calculate confidence score based on document analysis"""
        confidence = 100
        issues = []
        warnings = []
        
        # Check for Sharedo elements
        has_sharedo = (
            len(analysis["sharedo_elements"]["content_controls"]) > 0 or
            len(analysis["sharedo_elements"]["placeholders"]) > 0
        )
        
        if not has_sharedo:
            warnings.append("No Sharedo elements detected - may be a regular document")
            confidence -= 10
        
        # Complexity deductions
        if analysis["complexity_indicators"]["has_nested_tables"]:
            issues.append("Contains nested tables - may require manual review")
            confidence -= 15
        
        if analysis["complexity_indicators"]["has_custom_styles"]:
            warnings.append("Uses custom Word styles - formatting may vary")
            confidence -= 5
        
        if analysis["complexity_indicators"]["has_macros"]:
            issues.append("Contains macros - cannot be converted")
            confidence -= 20
        
        # Check for unsupported conditionals
        for conditional in analysis["sharedo_elements"]["conditionals"]:
            if "#foreach" in conditional:
                warnings.append("Contains #foreach loops - verify output")
                confidence -= 5
                break
        
        # Large document warning
        if analysis["statistics"]["total_words"] > 5000:
            warnings.append(f"Large document ({analysis['statistics']['total_words']} words) - verify performance")
            confidence -= 5
        
        # Many tables warning
        if analysis["statistics"]["tables"] > 10:
            warnings.append(f"Many tables ({analysis['statistics']['tables']}) - verify layout")
            confidence -= 5
        
        # Ensure confidence doesn't go below 0
        confidence = max(0, confidence)
        
        return confidence, issues, warnings
    
    def convert_document(self, docx_path, analysis):
        """Convert document to HTML based on analysis"""
        doc = Document(docx_path)
        html_parts = []
        
        # HTML header
        html_parts.append(self._get_html_header(docx_path.stem))
        
        # Process content based on whether it's a Sharedo template
        has_sharedo = len(analysis["sharedo_elements"]["content_controls"]) > 0
        
        if has_sharedo:
            # Process as Sharedo template
            html_parts.append(self._process_sharedo_content(doc, analysis))
        else:
            # Process as regular document
            html_parts.append(self._process_regular_content(doc))
        
        # HTML footer
        html_parts.append(self._get_html_footer())
        
        # Beautify
        soup = BeautifulSoup(''.join(html_parts), 'html.parser')
        return soup.prettify()
    
    def _get_html_header(self, title):
        """Generate HTML header"""
        return f'''<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        body {{
            font-family: 'Calibri', Arial, sans-serif;
            font-size: 11pt;
            line-height: 1.5;
            color: #000000;
            margin: 0;
            padding: 0;
            background-color: #f5f5f5;
        }}
        .document-container {{
            max-width: 816px;
            margin: 0 auto;
            background-color: #ffffff;
            padding: 72px 90px;
            box-shadow: 0 0 10px rgba(0,0,0,0.1);
        }}
        p {{ margin: 0 0 6pt 0; text-align: left; }}
        h1, h2, h3 {{ margin: 12pt 0 6pt 0; font-weight: bold; }}
        [data-tag] {{
            background: #e6f7ff;
            padding: 1px 3px;
            border-radius: 2px;
            font-family: monospace;
            font-size: 10pt;
        }}
        [data-content-block] {{
            background-color: #f8f9fa;
            padding: 8px 12px;
            margin: 12pt 0;
            border-left: 3px solid #dee2e6;
        }}
        [data-if] {{
            background-color: rgba(255, 243, 205, 0.3);
            border: 1px solid #ffc107;
            padding: 10px;
            margin: 12pt 0;
        }}
        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 12pt 0;
        }}
        th, td {{
            padding: 6pt;
            border: 1px solid #dee2e6;
            text-align: left;
        }}
    </style>
</head>
<body>
<div class="document-container">'''
    
    def _get_html_footer(self):
        """Generate HTML footer"""
        return '''
</div>
</body>
</html>'''
    
    def _process_sharedo_content(self, doc, analysis):
        """Process document as Sharedo template"""
        html_parts = []
        
        # Map content controls to their positions
        control_map = {c["tag"]: c for c in analysis["sharedo_elements"]["content_controls"]}
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Check for content blocks
            for block_tag in analysis["sharedo_elements"]["content_blocks"]:
                if block_tag in control_map:
                    control = control_map[block_tag]
                    if control.get("alias") and "Sharedo ContentBlock" in text:
                        html_parts.append(f'<div data-content-block="{block_tag}"><p>{control["alias"]}</p></div>')
                        text = ""
                        break
            
            # Check for sections
            for section_tag in analysis["sharedo_elements"]["sections"]:
                if section_tag in control_map and section_tag in text:
                    html_parts.append(f'<div data-section="{section_tag}">')
                    # Process section content
                    html_parts.append(f'<p>{text}</p>')
                    html_parts.append('</div>')
                    text = ""
                    break
            
            # Process regular paragraph with tag replacement
            if text:
                # Replace Sharedo tags
                for tag in analysis["sharedo_elements"]["tags"]:
                    if "Sharedo Tag:" in text:
                        text = re.sub(
                            r'Sharedo Tag:\s*' + re.escape(tag),
                            f'<span data-tag="{tag}">{tag}</span>',
                            text
                        )
                
                # Replace placeholders
                text = re.sub(r'\[_+\]', '<span data-tag="placeholder">[_____]</span>', text)
                
                # Apply formatting
                if para.style and para.style.name.startswith('Heading'):
                    level = re.search(r'\d', para.style.name)
                    level = level.group() if level else '2'
                    html_parts.append(f'<h{level}>{text}</h{level}>')
                else:
                    # Check for bold/italic in runs
                    if para.runs:
                        formatted_text = ""
                        for run in para.runs:
                            run_text = run.text
                            if run.bold:
                                run_text = f'<strong>{run_text}</strong>'
                            if run.italic:
                                run_text = f'<em>{run_text}</em>'
                            formatted_text += run_text
                        html_parts.append(f'<p>{formatted_text}</p>')
                    else:
                        html_parts.append(f'<p>{text}</p>')
        
        # Process tables
        for table in doc.tables:
            html_parts.append('<table>')
            for row_idx, row in enumerate(table.rows):
                if row_idx == 0:
                    html_parts.append('<thead><tr>')
                    for cell in row.cells:
                        html_parts.append(f'<th>{cell.text}</th>')
                    html_parts.append('</tr></thead>')
                else:
                    if row_idx == 1:
                        html_parts.append('<tbody>')
                    html_parts.append('<tr>')
                    for cell in row.cells:
                        html_parts.append(f'<td>{cell.text}</td>')
                    html_parts.append('</tr>')
            if len(table.rows) > 1:
                html_parts.append('</tbody>')
            html_parts.append('</table>')
        
        return '\n'.join(html_parts)
    
    def _process_regular_content(self, doc):
        """Process document as regular content"""
        html_parts = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Replace placeholders if any
            text = re.sub(r'\[_+\]', '<span class="placeholder">[_____]</span>', text)
            
            # Apply paragraph formatting
            if para.style and para.style.name.startswith('Heading'):
                level = re.search(r'\d', para.style.name)
                level = level.group() if level else '2'
                html_parts.append(f'<h{level}>{text}</h{level}>')
            else:
                # Check for formatting in runs
                if para.runs:
                    formatted_text = ""
                    for run in para.runs:
                        run_text = run.text
                        if run.bold:
                            run_text = f'<strong>{run_text}</strong>'
                        if run.italic:
                            run_text = f'<em>{run_text}</em>'
                        if run.underline:
                            run_text = f'<u>{run_text}</u>'
                        formatted_text += run_text
                    html_parts.append(f'<p>{formatted_text}</p>')
                else:
                    html_parts.append(f'<p>{text}</p>')
        
        # Process tables
        for table in doc.tables:
            html_parts.append('<table>')
            for row_idx, row in enumerate(table.rows):
                html_parts.append('<tr>')
                for cell in row.cells:
                    tag = 'th' if row_idx == 0 else 'td'
                    html_parts.append(f'<{tag}>{cell.text}</{tag}>')
                html_parts.append('</tr>')
            html_parts.append('</table>')
        
        return '\n'.join(html_parts)
    
    def generate_final_report(self):
        """Generate comprehensive HTML report"""
        report_path = self.output_folder / "conversion_report.html"
        
        # Sort files by confidence for easier review
        files_to_review = [f for f in self.report_data["files"] if f.get("requires_review", False)]
        files_successful = [f for f in self.report_data["files"] if f["status"] == "success" and not f.get("requires_review", False)]
        files_failed = [f for f in self.report_data["files"] if f["status"] == "failed"]
        
        html_content = f'''<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Sharedo Conversion Report</title>
    <style>
        body {{
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 1200px;
            margin: 0 auto;
            padding: 20px;
            background: #f5f5f5;
        }}
        h1, h2, h3 {{ color: #2c3e50; }}
        .header {{
            background: white;
            padding: 20px;
            border-radius: 8px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
            margin-bottom: 20px;
        }}
        .summary {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
            gap: 20px;
            margin: 20px 0;
        }}
        .stat-card {{
            background: white;
            padding: 15px;
            border-radius: 8px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
            text-align: center;
        }}
        .stat-value {{
            font-size: 2em;
            font-weight: bold;
            color: #3498db;
        }}
        .stat-label {{
            color: #7f8c8d;
            font-size: 0.9em;
        }}
        .file-section {{
            background: white;
            padding: 20px;
            border-radius: 8px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
            margin-bottom: 20px;
        }}
        .file-card {{
            border: 1px solid #e0e0e0;
            padding: 15px;
            margin: 10px 0;
            border-radius: 4px;
            background: #fafafa;
        }}
        .file-card.success {{ border-left: 4px solid #27ae60; }}
        .file-card.warning {{ border-left: 4px solid #f39c12; }}
        .file-card.error {{ border-left: 4px solid #e74c3c; }}
        .confidence {{
            display: inline-block;
            padding: 3px 8px;
            border-radius: 4px;
            font-weight: bold;
        }}
        .confidence.high {{ background: #d4edda; color: #155724; }}
        .confidence.medium {{ background: #fff3cd; color: #856404; }}
        .confidence.low {{ background: #f8d7da; color: #721c24; }}
        .issue {{ color: #e74c3c; margin: 5px 0; }}
        .warning {{ color: #f39c12; margin: 5px 0; }}
        .tag {{ 
            display: inline-block;
            background: #e6f7ff;
            padding: 2px 6px;
            border-radius: 3px;
            font-size: 0.85em;
            margin: 2px;
            font-family: monospace;
        }}
        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 15px 0;
        }}
        th, td {{
            padding: 10px;
            text-align: left;
            border-bottom: 1px solid #e0e0e0;
        }}
        th {{
            background: #f8f9fa;
            font-weight: bold;
        }}
        .timestamp {{
            color: #7f8c8d;
            font-size: 0.9em;
        }}
    </style>
</head>
<body>
    <div class="header">
        <h1>📊 Sharedo Batch Conversion Report</h1>
        <p class="timestamp">Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}</p>
    </div>
    
    <div class="summary">
        <div class="stat-card">
            <div class="stat-value">{self.report_data['total_files']}</div>
            <div class="stat-label">Total Files</div>
        </div>
        <div class="stat-card">
            <div class="stat-value">{self.report_data['successful']}</div>
            <div class="stat-label">Successful</div>
        </div>
        <div class="stat-card">
            <div class="stat-value">{len(files_to_review)}</div>
            <div class="stat-label">Need Review</div>
        </div>
        <div class="stat-card">
            <div class="stat-value">{self.report_data['failed']}</div>
            <div class="stat-label">Failed</div>
        </div>
    </div>
'''
        
        # Files requiring review
        if files_to_review:
            html_content += '''
    <div class="file-section">
        <h2>⚠️ Files Requiring Review</h2>
        <p>These files have lower confidence scores or issues that require manual verification.</p>
'''
            for file in sorted(files_to_review, key=lambda x: x.get('confidence_score', 0)):
                confidence_class = 'low' if file['confidence_score'] < 70 else 'medium'
                html_content += f'''
        <div class="file-card warning">
            <h3>{file['filename']} 
                <span class="confidence {confidence_class}">
                    {file['confidence_score']}% confidence
                </span>
            </h3>
'''
                if file.get('issues'):
                    html_content += '<h4>Issues:</h4>'
                    for issue in file['issues']:
                        html_content += f'<div class="issue">❌ {issue}</div>'
                
                if file.get('warnings'):
                    html_content += '<h4>Warnings:</h4>'
                    for warning in file['warnings']:
                        html_content += f'<div class="warning">⚠️ {warning}</div>'
                
                if file.get('sharedo_elements'):
                    html_content += '<h4>Detected Elements:</h4>'
                    html_content += '<div>'
                    for element_type, elements in file['sharedo_elements'].items():
                        if elements:
                            html_content += f'<strong>{element_type}:</strong> {len(elements)} '
                    html_content += '</div>'
                
                html_content += '''
        </div>
'''
            html_content += '''
    </div>
'''
        
        # Successful files
        if files_successful:
            html_content += '''
    <div class="file-section">
        <h2>✅ Successfully Converted</h2>
        <table>
            <thead>
                <tr>
                    <th>Filename</th>
                    <th>Confidence</th>
                    <th>Sharedo Elements</th>
                    <th>Statistics</th>
                </tr>
            </thead>
            <tbody>
'''
            for file in files_successful:
                elements_count = sum(len(v) for v in file.get('sharedo_elements', {}).values() if isinstance(v, list))
                stats = file.get('statistics', {})
                html_content += f'''
                <tr>
                    <td>{file['filename']}</td>
                    <td><span class="confidence high">{file['confidence_score']}%</span></td>
                    <td>{elements_count} elements</td>
                    <td>{stats.get('paragraphs', 0)} paragraphs, {stats.get('tables', 0)} tables</td>
                </tr>
'''
            html_content += '''
            </tbody>
        </table>
    </div>
'''
        
        # Failed files
        if files_failed:
            html_content += '''
    <div class="file-section">
        <h2>❌ Failed Conversions</h2>
'''
            for file in files_failed:
                html_content += f'''
        <div class="file-card error">
            <h3>{file['filename']}</h3>
            <p><strong>Error:</strong> {file.get('error', 'Unknown error')}</p>
            <details>
                <summary>View traceback</summary>
                <pre>{file.get('traceback', 'No traceback available')}</pre>
            </details>
        </div>
'''
            html_content += '''
    </div>
'''
        
        # Recommendations
        html_content += '''
    <div class="file-section">
        <h2>📝 Recommendations</h2>
        <ul>
            <li><strong>Files with &lt;90% confidence:</strong> Review the generated HTML to ensure all Sharedo tags are properly preserved</li>
            <li><strong>Nested tables:</strong> Verify table layout renders correctly in email clients</li>
            <li><strong>Custom styles:</strong> Check that formatting matches the original document</li>
            <li><strong>Complex conditionals:</strong> Test conditional logic with sample data</li>
        </ul>
    </div>
</body>
</html>'''
        
        # Save report
        report_path.write_text(html_content, encoding='utf-8')
        
        # Also save JSON version for programmatic access
        json_report_path = self.output_folder / "conversion_report.json"
        with open(json_report_path, 'w', encoding='utf-8') as f:
            json.dump(self.report_data, f, indent=2, default=str)


def main():
    """Main entry point for batch converter"""
    converter = SharedoBatchConverter(
        input_folder="Input",
        output_folder="Output"
    )
    
    converter.process_all_documents()


if __name__ == "__main__":
    main()