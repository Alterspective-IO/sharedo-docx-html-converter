#!/usr/bin/env python3
"""
Corrected Sharedo DOCX to HTML Converter
Properly replaces Sharedo placeholders with data-tag elements
"""

import re
import json
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
from docx import Document
from bs4 import BeautifulSoup
import html as html_module

class SharedoCorrectConverter:
    """Corrected converter that properly handles Sharedo content controls"""
    
    def __init__(self):
        self.content_controls = []
        self.tag_mapping = {}
        
    def load_metadata(self, json_path='word_tags.json'):
        """Load extracted Word metadata"""
        if Path(json_path).exists():
            with open(json_path, 'r') as f:
                data = json.load(f)
                self.content_controls = data.get('content_controls', [])
                
                # Build mapping of text to tags
                for control in self.content_controls:
                    # Map the display text to the tag
                    if control['text']:
                        # Clean up the text for matching
                        clean_text = control['text'].replace('Sharedo Tag:', '').strip()
                        clean_text = re.sub(r'\s+', ' ', clean_text)  # Normalize spaces
                        self.tag_mapping[clean_text] = control['tag']
                        
                        # Also map the original text
                        self.tag_mapping[control['text']] = control['tag']
                        
                        # Map partial matches for tags embedded in text
                        if 'Sharedo Tag:' in control['text']:
                            tag_only = control['tag']
                            self.tag_mapping[tag_only] = tag_only
                
                print(f"Loaded {len(self.content_controls)} content controls")
                print(f"Created {len(self.tag_mapping)} tag mappings")
        
        return self.content_controls
    
    def convert(self, docx_path, output_path=None):
        """Convert DOCX to Sharedo HTML"""
        
        # Load metadata
        self.load_metadata()
        
        # Process document
        doc = Document(docx_path)
        html_content = self._generate_html(doc)
        
        if output_path:
            Path(output_path).write_text(html_content, encoding='utf-8')
            print(f"✅ HTML saved to: {output_path}")
        
        return html_content
    
    def _generate_html(self, doc):
        """Generate HTML with Sharedo tags"""
        html_parts = []
        
        # Header
        html_parts.append(self._get_html_header())
        
        # Body
        body_content = self._process_document(doc)
        html_parts.append(body_content)
        
        # Footer
        html_parts.append(self._get_html_footer())
        
        # Beautify
        soup = BeautifulSoup(''.join(html_parts), 'html.parser')
        return soup.prettify()
    
    def _get_html_header(self):
        """HTML header matching example template style"""
        return '''<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Sharedo Email Template</title>
    <style type="text/css">
        body {
            font-family: Arial, Helvetica, sans-serif;
            font-size: 14px;
            line-height: 1.6;
            color: #333;
            margin: 20px;
        }
        p { margin: 0 0 10px 0; }
        .table { width: 100%; border-collapse: collapse; margin: 15px 0; }
        .table th, .table td { padding: 8px; border: 1px solid #ddd; text-align: left; }
    </style>
</head>
<body>'''
    
    def _get_html_footer(self):
        return '''</body>
</html>'''
    
    def _process_document(self, doc):
        """Process document and replace Sharedo placeholders"""
        content_parts = []
        current_section = None
        section_buffer = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                content_parts.append('<p>&nbsp;</p>')
                continue
            
            # Check for content blocks
            if 'Sharedo ContentBlock:' in text:
                # Extract the content block tag
                for control in self.content_controls:
                    if control['text'] in text:
                        tag = control['tag']
                        alias = control['alias'].replace('Sharedo ContentBlock: ', '')
                        html = f'<div data-content-block="{html_module.escape(tag)}">\n'
                        html += f'    <p>{html_module.escape(alias)}</p>\n'
                        html += '</div>'
                        content_parts.append(html)
                        break
                continue
            
            # Check for sections
            section_found = False
            for control in self.content_controls:
                if control.get('alias', '').startswith('Sharedo Section:') and control['text'] in text:
                    # Close previous section if exists
                    if current_section:
                        content_parts.append(f'<div data-section="{html_module.escape(current_section)}">')
                        content_parts.extend(section_buffer)
                        content_parts.append('</div>')
                        section_buffer = []
                    
                    # Start new section
                    current_section = control['tag']
                    # Process the section content
                    section_text = control['text']
                    section_html = self._replace_tags_in_text(section_text, para)
                    section_buffer.append(section_html)
                    section_found = True
                    break
            
            if section_found:
                continue
            
            # Regular paragraph - replace any Sharedo tags
            para_html = self._replace_tags_in_text(text, para)
            
            if current_section:
                section_buffer.append(para_html)
            else:
                content_parts.append(para_html)
        
        # Close any open section
        if current_section:
            content_parts.append(f'<div data-section="{html_module.escape(current_section)}">')
            content_parts.extend(section_buffer)
            content_parts.append('</div>')
        
        # Process tables
        for table in doc.tables:
            table_html = self._process_table(table)
            content_parts.append(table_html)
        
        return '\n'.join(content_parts)
    
    def _replace_tags_in_text(self, text, paragraph=None):
        """Replace Sharedo placeholders with proper tags"""
        original_text = text
        
        # Replace each content control found in the text
        for control in self.content_controls:
            if 'Sharedo Tag:' in control['text']:
                # Multiple possible patterns to match
                patterns_to_try = [
                    control['text'],  # Full text
                    control['text'].replace('Sharedo Tag:', '').strip(),  # Without prefix
                    re.sub(r'\s+', ' ', control['text']),  # Normalized spaces
                    control['tag'],  # Just the tag itself
                ]
                
                for pattern in patterns_to_try:
                    if pattern in text:
                        replacement = f'<span data-tag="{html_module.escape(control["tag"])}">{html_module.escape(control["tag"])}</span>'
                        text = text.replace(pattern, replacement)
                        break
        
        # Apply paragraph formatting if provided
        if paragraph:
            # Check for bold, italic, underline in runs
            formatted_parts = []
            if paragraph.runs:
                for run in paragraph.runs:
                    run_text = run.text
                    
                    # Replace tags in run text
                    for control in self.content_controls:
                        if 'Sharedo Tag:' in control.get('text', ''):
                            patterns = [
                                control['text'],
                                control['text'].replace('Sharedo Tag:', '').strip(),
                                control['tag']
                            ]
                            for pattern in patterns:
                                if pattern in run_text:
                                    run_text = run_text.replace(
                                        pattern,
                                        f'<span data-tag="{html_module.escape(control["tag"])}">{html_module.escape(control["tag"])}</span>'
                                    )
                    
                    # Apply formatting
                    if run.bold:
                        run_text = f'<strong>{run_text}</strong>'
                    if run.italic:
                        run_text = f'<em>{run_text}</em>'
                    if run.underline:
                        run_text = f'<u>{run_text}</u>'
                    
                    formatted_parts.append(run_text)
                
                if formatted_parts:
                    text = ''.join(formatted_parts)
            
            # Check for heading
            if paragraph.style and paragraph.style.name.startswith('Heading'):
                level = self._get_heading_level(paragraph.style.name)
                return f'<h{level}>{text}</h{level}>'
        
        return f'<p>{text}</p>'
    
    def _process_table(self, table):
        """Process table with Sharedo tag replacement"""
        html_parts = ['<figure class="table">\n<table>']
        
        if table.rows:
            # Header row
            html_parts.append('\n<thead>\n<tr>')
            for cell in table.rows[0].cells:
                cell_text = cell.text.strip()
                # Replace Sharedo tags
                for control in self.content_controls:
                    if 'Sharedo Tag:' in control.get('text', '') and control['text'] in cell_text:
                        cell_text = cell_text.replace(
                            control['text'],
                            f'<span data-tag="{html_module.escape(control["tag"])}">{html_module.escape(control["tag"])}</span>'
                        )
                html_parts.append(f'\n<th>{cell_text}</th>')
            html_parts.append('\n</tr>\n</thead>')
            
            # Body rows
            if len(table.rows) > 1:
                html_parts.append('\n<tbody>')
                for row in table.rows[1:]:
                    html_parts.append('\n<tr>')
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        # Replace Sharedo tags
                        for control in self.content_controls:
                            if 'Sharedo Tag:' in control.get('text', '') and control['text'] in cell_text:
                                cell_text = cell_text.replace(
                                    control['text'],
                                    f'<span data-tag="{html_module.escape(control["tag"])}">{html_module.escape(control["tag"])}</span>'
                                )
                        html_parts.append(f'\n<td>{cell_text}</td>')
                    html_parts.append('\n</tr>')
                html_parts.append('\n</tbody>')
        
        html_parts.append('\n</table>\n</figure>')
        return ''.join(html_parts)
    
    def _get_heading_level(self, style_name):
        """Extract heading level"""
        try:
            return int(re.search(r'\d', style_name).group())
        except:
            return 2


def main():
    """Convert SUPLC1031.docx with proper Sharedo tag handling"""
    converter = SharedoCorrectConverter()
    
    docx_file = "SUPLC1031.docx"
    output_file = "SUPLC1031_sharedo_corrected.html"
    
    print("🚀 Starting Corrected Sharedo Conversion")
    print("=" * 50)
    
    # Convert
    html_content = converter.convert(docx_file, output_file)
    
    # Validate
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Count Sharedo elements
    data_tags = soup.find_all(attrs={'data-tag': True})
    content_blocks = soup.find_all(attrs={'data-content-block': True})
    sections = soup.find_all(attrs={'data-section': True})
    
    print("\n📊 Conversion Results:")
    print(f"  • Sharedo Tags: {len(data_tags)}")
    if data_tags:
        print("  • Sample tags found:")
        for tag in data_tags[:5]:
            print(f"    - {tag.get('data-tag')}")
    
    print(f"  • Content Blocks: {len(content_blocks)}")
    if content_blocks:
        for block in content_blocks:
            print(f"    - {block.get('data-content-block')}")
    
    print(f"  • Sections: {len(sections)}")
    if sections:
        for section in sections:
            print(f"    - {section.get('data-section')}")
    
    print("\n✅ Conversion complete!")
    
    return html_content


if __name__ == "__main__":
    main()