#!/usr/bin/env python3
"""Analyze DOCX file to understand structure and special tags"""

import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import RGBColor
    import re
except ImportError:
    print("Installing required packages...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import RGBColor
    import re

def analyze_docx(file_path):
    """Analyze DOCX document for structure and special tags"""
    doc = Document(file_path)
    
    print("=" * 60)
    print(f"DOCUMENT ANALYSIS: {file_path}")
    print("=" * 60)
    
    # Analyze paragraphs
    print("\n📄 PARAGRAPHS OVERVIEW:")
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    
    # Look for special patterns (If conditions, tags, etc.)
    special_patterns = {
        'if_conditions': re.compile(r'\{\{#if\s+.*?\}\}|\{\{/if\}\}', re.IGNORECASE),
        'variables': re.compile(r'\{\{[^#/].*?\}\}'),
        'loops': re.compile(r'\{\{#each\s+.*?\}\}|\{\{/each\}\}', re.IGNORECASE),
        'tags': re.compile(r'<[^>]+>'),
        'merge_fields': re.compile(r'\[\[.*?\]\]'),
    }
    
    pattern_counts = {name: 0 for name in special_patterns}
    sample_content = []
    
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text:
            # Check for special patterns
            for pattern_name, pattern in special_patterns.items():
                matches = pattern.findall(text)
                if matches:
                    pattern_counts[pattern_name] += len(matches)
                    if len(sample_content) < 5:  # Collect first 5 samples
                        sample_content.append(f"Para {i+1}: {text[:150]}...")
            
    print("\n🔍 SPECIAL PATTERNS FOUND:")
    for pattern_name, count in pattern_counts.items():
        if count > 0:
            print(f"  • {pattern_name.replace('_', ' ').title()}: {count} occurrences")
    
    print("\n📝 SAMPLE CONTENT WITH SPECIAL PATTERNS:")
    for sample in sample_content[:5]:
        print(f"  {sample}")
    
    # Analyze tables
    if doc.tables:
        print(f"\n📊 TABLES: {len(doc.tables)} found")
        for i, table in enumerate(doc.tables):
            print(f"  • Table {i+1}: {len(table.rows)} rows x {len(table.columns)} columns")
    
    # Analyze styles
    print("\n🎨 FORMATTING ANALYSIS:")
    styles_used = set()
    fonts_used = set()
    
    for paragraph in doc.paragraphs:
        if paragraph.style.name:
            styles_used.add(paragraph.style.name)
        for run in paragraph.runs:
            if run.font.name:
                fonts_used.add(run.font.name)
    
    print(f"  • Unique styles: {len(styles_used)}")
    if styles_used:
        print(f"    Samples: {', '.join(list(styles_used)[:5])}")
    print(f"  • Unique fonts: {len(fonts_used)}")
    if fonts_used:
        print(f"    Fonts: {', '.join(fonts_used)}")
    
    # Full content preview
    print("\n📖 FULL CONTENT PREVIEW (first 500 chars):")
    full_text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
    print(full_text[:500] + "...")
    
    return doc

if __name__ == "__main__":
    file_path = "We refer to the telephone conversation.docx"
    if Path(file_path).exists():
        analyze_docx(file_path)
    else:
        print(f"Error: File '{file_path}' not found!")