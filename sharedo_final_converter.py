#!/usr/bin/env python3
"""
Final Sharedo DOCX to HTML Converter
Handles all Sharedo content controls, sections, and tags
"""

import re
import json
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup
import html as html_module

class SharedoFinalConverter:
    """Complete Sharedo DOCX to HTML converter with full content control support"""
    
    def __init__(self):
        self.content_controls = []
        self.sections = {}
        self.tags = {}
        
    def extract_sharedo_metadata(self, docx_path):
        """Extract all Sharedo metadata from the document"""
        metadata = {
            'content_controls': [],
            'sections': {},
            'content_blocks': {},
            'tags': {}
        }
        
        with zipfile.ZipFile(docx_path, 'r') as docx_zip:
            # Parse main document XML
            if 'word/document.xml' in docx_zip.namelist():
                doc_xml = docx_zip.read('word/document.xml').decode('utf-8')
                
                namespaces = {
                    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                }
                
                root = ET.fromstring(doc_xml)
                
                # Extract all content controls
                for sdt in root.findall('.//w:sdt', namespaces):
                    sdt_pr = sdt.find('w:sdtPr', namespaces)
                    if sdt_pr is not None:
                        tag_elem = sdt_pr.find('w:tag', namespaces)
                        alias_elem = sdt_pr.find('w:alias', namespaces)
                        
                        tag = tag_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') if tag_elem is not None else None
                        alias = alias_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') if alias_elem is not None else None
                        
                        # Get content text
                        sdt_content = sdt.find('.//w:sdtContent', namespaces)
                        text_content = ''
                        if sdt_content is not None:
                            texts = sdt_content.findall('.//w:t', namespaces)
                            text_content = ' '.join([t.text for t in texts if t.text])
                        
                        if tag:
                            control = {
                                'tag': tag,
                                'alias': alias,
                                'text': text_content.strip()
                            }
                            
                            # Categorize by type
                            if alias and 'ContentBlock' in alias:
                                metadata['content_blocks'][tag] = control
                            elif alias and 'Section' in alias:
                                metadata['sections'][tag] = control
                            elif alias and 'Tag' in alias:
                                metadata['tags'][tag] = control
                            
                            metadata['content_controls'].append(control)
        
        return metadata
    
    def convert(self, docx_path, output_path=None):
        """Convert DOCX to Sharedo HTML template"""
        
        # Extract Sharedo metadata
        metadata = self.extract_sharedo_metadata(docx_path)
        self.content_controls = metadata['content_controls']
        self.sections = metadata['sections']
        self.tags = metadata['tags']
        
        print(f"📊 Found Sharedo elements:")
        print(f"  • Content Controls: {len(self.content_controls)}")
        print(f"  • Sections: {len(self.sections)}")
        print(f"  • Tags: {len(self.tags)}")
        print(f"  • Content Blocks: {len(metadata['content_blocks'])}")
        
        # Process document
        doc = Document(docx_path)
        html_content = self._generate_html(doc, metadata)
        
        if output_path:
            Path(output_path).write_text(html_content, encoding='utf-8')
            print(f"✅ Sharedo HTML template saved to: {output_path}")
        
        return html_content
    
    def _generate_html(self, doc, metadata):
        """Generate Sharedo-compatible HTML"""
        html_parts = []
        
        # HTML header
        html_parts.append(self._get_html_header())
        
        # Process document body
        body_content = self._process_document(doc, metadata)
        html_parts.append(body_content)
        
        # HTML footer
        html_parts.append(self._get_html_footer())
        
        # Beautify
        soup = BeautifulSoup(''.join(html_parts), 'html.parser')
        return soup.prettify()
    
    def _get_html_header(self):
        """Sharedo email template HTML header"""
        return '''<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Sharedo Template</title>
    <style type="text/css">
        /* Email Reset */
        body, table, td, a { -webkit-text-size-adjust: 100%; -ms-text-size-adjust: 100%; }
        table, td { mso-table-lspace: 0pt; mso-table-rspace: 0pt; }
        
        /* Base Styles */
        body {
            margin: 0;
            padding: 0;
            font-family: Arial, Helvetica, sans-serif;
            font-size: 14px;
            line-height: 1.6;
            color: #333333;
        }
        
        p { margin: 0 0 10px 0; }
        h1, h2, h3 { margin: 0 0 10px 0; }
        
        /* Sharedo Elements */
        [data-tag] {
            background-color: #e6f7ff;
            padding: 2px 4px;
            border-radius: 2px;
            font-family: monospace;
            font-size: 0.9em;
        }
        
        [data-content-block] {
            background-color: #f0f0f0;
            padding: 10px;
            margin: 10px 0;
            border: 1px dashed #999;
        }
        
        [data-section] {
            border-left: 3px solid #0969da;
            padding-left: 10px;
            margin: 10px 0;
        }
        
        .table {
            width: 100%;
            border-collapse: collapse;
            margin: 15px 0;
        }
        
        .table th, .table td {
            padding: 8px;
            border: 1px solid #ddd;
            text-align: left;
        }
        
        .table th {
            background-color: #f8f9fa;
        }
    </style>
</head>
<body>'''
    
    def _get_html_footer(self):
        """HTML footer"""
        return '''</body>
</html>'''
    
    def _process_document(self, doc, metadata):
        """Process document with Sharedo content control preservation"""
        content_parts = []
        current_section = None
        section_content = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Check if this paragraph contains a content control
            control = self._find_matching_control(text, metadata)
            
            if control:
                control_type = self._get_control_type(control)
                
                if control_type == 'content_block':
                    # Content block
                    html = f'<div data-content-block="{html_module.escape(control["tag"])}">'
                    html += f'<p>* {control["alias"].replace("Sharedo ContentBlock: ", "")}</p>'
                    html += '</div>'
                    content_parts.append(html)
                    
                elif control_type == 'section':
                    # Section start/end
                    if current_section:
                        # Close previous section
                        content_parts.append(f'<div data-section="{html_module.escape(current_section)}">')
                        content_parts.extend(section_content)
                        content_parts.append('</div>')
                        section_content = []
                    
                    current_section = control['tag']
                    # Process section content
                    section_html = self._process_section_content(control['text'])
                    section_content.append(section_html)
                    
                elif control_type == 'tag':
                    # Sharedo tag
                    tag_name = control['tag']
                    html = f'<span data-tag="{html_module.escape(tag_name)}">{html_module.escape(tag_name)}</span>'
                    
                    # If it's part of a paragraph, wrap it
                    if current_section:
                        section_content.append(f'<p>{html}</p>')
                    else:
                        content_parts.append(f'<p>{html}</p>')
            else:
                # Regular paragraph
                para_html = self._process_paragraph(para)
                if current_section:
                    section_content.append(para_html)
                else:
                    content_parts.append(para_html)
        
        # Close any open section
        if current_section:
            content_parts.append(f'<div data-section="{html_module.escape(current_section)}">')
            content_parts.extend(section_content)
            content_parts.append('</div>')
        
        # Process tables
        for table in doc.tables:
            table_html = self._process_table(table)
            content_parts.append(table_html)
        
        return '\n'.join(content_parts)
    
    def _find_matching_control(self, text, metadata):
        """Find matching content control for given text"""
        for control in metadata['content_controls']:
            # Check if the text contains the control's display text
            if control['text'] and control['text'] in text:
                return control
            # Also check for alias matches
            if control['alias'] and control['alias'] in text:
                return control
        return None
    
    def _get_control_type(self, control):
        """Determine the type of content control"""
        alias = control.get('alias', '')
        if 'ContentBlock' in alias:
            return 'content_block'
        elif 'Section' in alias:
            return 'section'
        elif 'Tag' in alias:
            return 'tag'
        return 'unknown'
    
    def _process_section_content(self, text):
        """Process content within a section"""
        # Clean up the text
        text = text.strip()
        
        # Remove the Sharedo markers
        text = re.sub(r'Sharedo (Section|Tag|ContentBlock):\s*', '', text)
        
        # Convert to HTML paragraphs
        paragraphs = text.split('\n')
        html_parts = []
        
        for para in paragraphs:
            para = para.strip()
            if para:
                # Check for Sharedo tags within the text
                para = self._replace_sharedo_tags(para)
                html_parts.append(f'<p>{para}</p>')
        
        return '\n'.join(html_parts)
    
    def _replace_sharedo_tags(self, text):
        """Replace Sharedo tag references in text"""
        # Pattern for Sharedo tags
        tag_pattern = re.compile(r'(context\.[a-zA-Z0-9._\-!?&=]+|document\.[a-zA-Z0-9._\-!?&=]+)')
        
        def replace_tag(match):
            tag = match.group(1)
            return f'<span data-tag="{html_module.escape(tag)}">{html_module.escape(tag)}</span>'
        
        return tag_pattern.sub(replace_tag, text)
    
    def _process_paragraph(self, paragraph):
        """Process a regular paragraph"""
        text = paragraph.text.strip()
        if not text:
            return ''
        
        # Check for Sharedo tags in the text
        text = self._replace_sharedo_tags(text)
        
        # Apply formatting from runs
        if paragraph.runs:
            formatted_text = self._format_runs(paragraph.runs)
            if formatted_text:
                text = formatted_text
        
        # Check for heading style
        if paragraph.style and paragraph.style.name.startswith('Heading'):
            level = self._get_heading_level(paragraph.style.name)
            return f'<h{level}>{text}</h{level}>'
        
        return f'<p>{text}</p>'
    
    def _format_runs(self, runs):
        """Format text runs with bold, italic, etc."""
        formatted_parts = []
        
        for run in runs:
            if not run.text:
                continue
            
            text = run.text
            
            # Replace Sharedo tags
            text = self._replace_sharedo_tags(text)
            
            # Apply formatting
            if run.bold:
                text = f'<strong>{text}</strong>'
            if run.italic:
                text = f'<em>{text}</em>'
            if run.underline:
                text = f'<u>{text}</u>'
            
            formatted_parts.append(text)
        
        return ''.join(formatted_parts) if formatted_parts else None
    
    def _process_table(self, table):
        """Process table with Sharedo support"""
        html_parts = ['<figure class="table"><table>']
        
        # Check if first row is header
        if table.rows:
            # Header row
            html_parts.append('<thead><tr>')
            for cell in table.rows[0].cells:
                cell_text = cell.text.strip()
                cell_text = self._replace_sharedo_tags(cell_text)
                html_parts.append(f'<th>{cell_text}</th>')
            html_parts.append('</tr></thead>')
            
            # Body rows
            if len(table.rows) > 1:
                html_parts.append('<tbody>')
                for row in table.rows[1:]:
                    html_parts.append('<tr>')
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        cell_text = self._replace_sharedo_tags(cell_text)
                        html_parts.append(f'<td>{cell_text}</td>')
                    html_parts.append('</tr>')
                html_parts.append('</tbody>')
        
        html_parts.append('</table></figure>')
        return '\n'.join(html_parts)
    
    def _get_heading_level(self, style_name):
        """Extract heading level from style name"""
        try:
            return int(re.search(r'\d', style_name).group())
        except:
            return 2


def main():
    """Convert SUPLC1031.docx with complete Sharedo support"""
    converter = SharedoFinalConverter()
    
    docx_file = "SUPLC1031.docx"
    output_file = "SUPLC1031_complete_sharedo.html"
    
    print("🚀 Starting Final Sharedo DOCX to HTML Conversion")
    print("=" * 50)
    print(f"Input: {docx_file}")
    
    # Perform conversion
    html_content = converter.convert(docx_file, output_file)
    
    # Validate output
    soup = BeautifulSoup(html_content, 'html.parser')
    
    stats = {
        'data_tags': len(soup.find_all(attrs={'data-tag': True})),
        'content_blocks': len(soup.find_all(attrs={'data-content-block': True})),
        'sections': len(soup.find_all(attrs={'data-section': True})),
        'tables': len(soup.find_all('table'))
    }
    
    print("\n📊 Conversion Complete:")
    print(f"  • Output: {output_file}")
    print(f"  • Sharedo Tags: {stats['data_tags']}")
    print(f"  • Content Blocks: {stats['content_blocks']}")
    print(f"  • Sections: {stats['sections']}")
    print(f"  • Tables: {stats['tables']}")
    
    print("\n✅ Sharedo template ready for use!")
    
    return html_content


if __name__ == "__main__":
    main()