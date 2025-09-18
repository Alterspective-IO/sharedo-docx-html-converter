#!/usr/bin/env python3
"""
Detailed Analysis of Document Conversions
Compares original DOCX files with converted HTML to identify patterns
"""

import os
import json
import requests
from pathlib import Path
from docx import Document
from bs4 import BeautifulSoup
from collections import defaultdict
from typing import Dict, List, Tuple

class DetailedAnalyzer:
    def __init__(self, api_url: str):
        self.api_url = api_url
        self.patterns = {
            "working_well": defaultdict(list),
            "needs_improvement": defaultdict(list),
            "completely_missing": defaultdict(list)
        }
        
    def analyze_document_pair(self, docx_path: str) -> Dict:
        """Analyze a DOCX file and its HTML conversion"""
        # Read original DOCX
        doc = Document(docx_path)
        original_content = self.extract_docx_content(doc)
        
        # Convert via API
        with open(docx_path, 'rb') as f:
            response = requests.post(
                f"{self.api_url}/api/v1/convert",
                files={'file': (os.path.basename(docx_path), f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
            )
        
        if response.status_code != 200:
            return {"error": f"Conversion failed: {response.status_code}"}
            
        result = response.json()
        soup = BeautifulSoup(result['html_content'], 'html.parser')
        
        # Detailed comparison
        analysis = {
            "document": os.path.basename(docx_path),
            "original": original_content,
            "converted": self.extract_html_content(soup),
            "sharedo_elements": result.get('sharedo_elements', {}),
            "confidence": result.get('confidence_score', 0),
            "issues": []
        }
        
        # Check specific patterns
        self.check_tag_preservation(original_content, soup, analysis)
        self.check_conditional_blocks(original_content, soup, analysis)
        self.check_formatting(doc, soup, analysis)
        self.check_tables(doc, soup, analysis)
        
        return analysis
    
    def extract_docx_content(self, doc) -> Dict:
        """Extract detailed content from DOCX"""
        content = {
            "paragraphs": [],
            "tables": [],
            "tags": [],
            "conditionals": [],
            "sections": []
        }
        
        # Extract paragraphs and tags
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                content["paragraphs"].append(text)
                
                # Look for Sharedo tags
                if "{{" in text or "{%" in text:
                    content["tags"].append(text)
                elif any(tag in text for tag in ["context.", "document.", "env."]):
                    content["tags"].append(text)
                    
                # Look for conditional markers
                if any(marker in text.lower() for marker in ["if ", "then", "else", "endif"]):
                    content["conditionals"].append(text)
        
        # Extract tables
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            content["tables"].append(table_data)
            
        return content
    
    def extract_html_content(self, soup) -> Dict:
        """Extract content from HTML"""
        content = {
            "paragraphs": [],
            "tables": [],
            "data_attributes": [],
            "sharedo_elements": []
        }
        
        # Extract paragraphs
        for p in soup.find_all('p'):
            text = p.get_text(strip=True)
            if text:
                content["paragraphs"].append(text)
        
        # Extract data attributes (Sharedo elements)
        for elem in soup.find_all(attrs={"data-tag": True}):
            content["data_attributes"].append({
                "tag": elem.get("data-tag"),
                "text": elem.get_text(strip=True)
            })
            
        for elem in soup.find_all(attrs={"data-section": True}):
            content["sharedo_elements"].append({
                "type": "section",
                "value": elem.get("data-section")
            })
            
        for elem in soup.find_all(attrs={"data-if": True}):
            content["sharedo_elements"].append({
                "type": "conditional",
                "value": elem.get("data-if")
            })
        
        # Extract tables
        for table in soup.find_all('table'):
            table_data = []
            for row in table.find_all('tr'):
                cells = row.find_all(['td', 'th'])
                row_data = [cell.get_text(strip=True) for cell in cells]
                table_data.append(row_data)
            content["tables"].append(table_data)
            
        return content
    
    def check_tag_preservation(self, original, soup, analysis):
        """Check how well tags are preserved"""
        original_tags = set(original.get("tags", []))
        converted_tags = set()
        
        # Find all elements with data-tag attribute
        for elem in soup.find_all(attrs={"data-tag": True}):
            converted_tags.add(elem.get("data-tag"))
        
        if original_tags:
            preserved = len(converted_tags) / len(original_tags) * 100
            analysis["tag_preservation"] = {
                "original_count": len(original_tags),
                "converted_count": len(converted_tags),
                "preservation_rate": preserved
            }
            
            if preserved >= 80:
                self.patterns["working_well"]["tags"].append(analysis["document"])
            elif preserved >= 50:
                self.patterns["needs_improvement"]["tags"].append(analysis["document"])
            else:
                self.patterns["completely_missing"]["tags"].append(analysis["document"])
                
            missing_tags = original_tags - converted_tags
            if missing_tags:
                analysis["issues"].append(f"Missing tags: {list(missing_tags)[:5]}")
    
    def check_conditional_blocks(self, original, soup, analysis):
        """Check conditional block handling"""
        original_conditionals = original.get("conditionals", [])
        converted_conditionals = soup.find_all(attrs={"data-if": True})
        
        if original_conditionals:
            if converted_conditionals:
                self.patterns["working_well"]["conditionals"].append(analysis["document"])
            else:
                self.patterns["completely_missing"]["conditionals"].append(analysis["document"])
                analysis["issues"].append("Conditional blocks not preserved")
    
    def check_formatting(self, doc, soup, analysis):
        """Check formatting preservation"""
        # Check for bold, italic, underline
        formatting_preserved = True
        
        for para in doc.paragraphs:
            for run in para.runs:
                if run.bold or run.italic or run.underline:
                    # Check if similar formatting exists in HTML
                    if not (soup.find('strong') or soup.find('b') or 
                           soup.find('em') or soup.find('i') or
                           soup.find('u')):
                        formatting_preserved = False
                        break
        
        if formatting_preserved:
            self.patterns["working_well"]["formatting"].append(analysis["document"])
        else:
            self.patterns["needs_improvement"]["formatting"].append(analysis["document"])
    
    def check_tables(self, doc, soup, analysis):
        """Check table preservation"""
        original_tables = len(doc.tables)
        converted_tables = len(soup.find_all('table'))
        
        if original_tables > 0:
            if original_tables == converted_tables:
                self.patterns["working_well"]["tables"].append(analysis["document"])
                analysis["table_preservation"] = "Perfect"
            elif converted_tables > 0:
                self.patterns["needs_improvement"]["tables"].append(analysis["document"])
                analysis["table_preservation"] = f"Partial ({converted_tables}/{original_tables})"
            else:
                self.patterns["completely_missing"]["tables"].append(analysis["document"])
                analysis["table_preservation"] = "Missing"
                analysis["issues"].append(f"Tables not preserved: {original_tables} expected")
    
    def generate_pattern_report(self) -> str:
        """Generate a report of identified patterns"""
        report = []
        report.append("=" * 80)
        report.append("DETAILED PATTERN ANALYSIS")
        report.append("=" * 80)
        report.append("")
        
        # What's working well
        report.append("✅ WORKING WELL")
        report.append("-" * 40)
        for category, docs in self.patterns["working_well"].items():
            if docs:
                report.append(f"\n{category.upper()} (Working in {len(docs)} documents):")
                for doc in docs[:3]:  # Show first 3 examples
                    report.append(f"  • {doc}")
        
        # What needs improvement
        report.append("\n" + "⚠️  NEEDS IMPROVEMENT")
        report.append("-" * 40)
        for category, docs in self.patterns["needs_improvement"].items():
            if docs:
                report.append(f"\n{category.upper()} (Issues in {len(docs)} documents):")
                for doc in docs[:3]:
                    report.append(f"  • {doc}")
        
        # What's completely missing
        report.append("\n" + "❌ COMPLETELY MISSING")
        report.append("-" * 40)
        for category, docs in self.patterns["completely_missing"].items():
            if docs:
                report.append(f"\n{category.upper()} (Missing in {len(docs)} documents):")
                for doc in docs[:3]:
                    report.append(f"  • {doc}")
        
        # Summary and recommendations
        report.append("\n" + "=" * 80)
        report.append("RECOMMENDATIONS")
        report.append("-" * 40)
        
        if self.patterns["completely_missing"]["tags"]:
            report.append("1. Priority: Fix tag extraction - tags are completely missing in some documents")
        if self.patterns["completely_missing"]["conditionals"]:
            report.append("2. Priority: Implement conditional block detection and conversion")
        if self.patterns["needs_improvement"]["tables"]:
            report.append("3. Improve table structure preservation")
        if self.patterns["needs_improvement"]["formatting"]:
            report.append("4. Enhance text formatting preservation (bold, italic, etc.)")
        
        return "\n".join(report)

def main():
    api_url = "https://sharedo-docx-converter-dev.politeground-d6f68ba9.australiaeast.azurecontainerapps.io"
    samples_dir = "/Users/igorsharedo/Documents/Samples"
    
    analyzer = DetailedAnalyzer(api_url)
    
    # Analyze a subset of documents for detailed patterns
    test_docs = [
        "Igors Agreement.docx",  # Simple - Grade C
        "AdminOperation_Print_Task_Template.docx",  # Complex - Grade C+
        "Draft Agreement copy.docx",  # Moderate - Grade A+
        "Acme Supermarkets - Certification.docx"  # Simple - Grade C
    ]
    
    print("Performing detailed analysis on selected documents...")
    print("=" * 80)
    
    detailed_results = []
    for doc_name in test_docs:
        doc_path = os.path.join(samples_dir, doc_name)
        if os.path.exists(doc_path):
            print(f"\nAnalyzing: {doc_name}")
            result = analyzer.analyze_document_pair(doc_path)
            detailed_results.append(result)
            
            # Print immediate findings
            if result.get("issues"):
                print(f"  Issues found: {len(result['issues'])}")
                for issue in result['issues'][:2]:
                    print(f"    - {issue}")
            else:
                print("  ✓ No major issues")
    
    # Generate pattern report
    pattern_report = analyzer.generate_pattern_report()
    print("\n" + pattern_report)
    
    # Save detailed results
    with open("detailed_analysis.json", "w") as f:
        json.dump({
            "results": detailed_results,
            "patterns": {
                "working_well": dict(analyzer.patterns["working_well"]),
                "needs_improvement": dict(analyzer.patterns["needs_improvement"]),
                "completely_missing": dict(analyzer.patterns["completely_missing"])
            }
        }, f, indent=2)
    
    print("\n" + "=" * 80)
    print("Detailed analysis complete!")
    print("Results saved to: detailed_analysis.json")

if __name__ == "__main__":
    main()