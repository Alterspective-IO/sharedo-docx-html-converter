#!/usr/bin/env python3
"""Debug script to see actual paragraph text in DOCX"""

from docx import Document

doc = Document("SUPLC1031.docx")

print("=" * 60)
print("ACTUAL PARAGRAPH TEXT IN DOCUMENT:")
print("=" * 60)

for i, para in enumerate(doc.paragraphs):
    text = para.text
    if text.strip():
        print(f"\nPara {i+1}:")
        print(f"  Text: [{text}]")
        
        # Show runs
        if para.runs:
            print(f"  Runs ({len(para.runs)}):")
            for j, run in enumerate(para.runs):
                if run.text:
                    print(f"    Run {j+1}: [{run.text}]")

print("\n" + "=" * 60)
print("TABLE CONTENT:")
print("=" * 60)

for t_idx, table in enumerate(doc.tables):
    print(f"\nTable {t_idx + 1}:")
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            if cell.text.strip():
                print(f"  Cell[{r_idx},{c_idx}]: [{cell.text}]")