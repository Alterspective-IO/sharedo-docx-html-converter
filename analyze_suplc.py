#!/usr/bin/env python3
"""Analyze SUPLC1031.docx for Sharedo tags and content"""

from pathlib import Path
from docx import Document
import re

def analyze_sharedo_doc(file_path):
    """Analyze document for Sharedo-specific patterns"""
    doc = Document(file_path)
    
    print("=" * 60)
    print(f"SHAREDO DOCUMENT ANALYSIS: {file_path}")
    print("=" * 60)
    
    # Sharedo-specific patterns
    sharedo_patterns = {
        'data_tags': re.compile(r'\{\{([^}]+)\}\}'),
        'conditionals': re.compile(r'#if\s+|#endif|#else'),
        'loops': re.compile(r'#foreach\s+|#endforeach'),
        'variables': re.compile(r'context\.[a-zA-Z.]+|document\.[a-zA-Z.]+'),
        'placeholders': re.compile(r'\[_+\]'),
        'merge_fields': re.compile(r'«([^»]+)»'),
    }
    
    findings = {name: [] for name in sharedo_patterns}
    full_text = []
    
    # Analyze paragraphs
    print(f"\n📄 DOCUMENT STRUCTURE:")
    print(f"  • Total paragraphs: {len(doc.paragraphs)}")
    print(f"  • Total tables: {len(doc.tables)}")
    
    # Extract and analyze text
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            full_text.append(text)
            # Check for patterns
            for pattern_name, pattern in sharedo_patterns.items():
                matches = pattern.findall(text)
                if matches:
                    findings[pattern_name].extend(matches)
                    if len(findings[pattern_name]) <= 3:  # Show first 3 examples
                        print(f"\n🔍 Found {pattern_name.replace('_', ' ').title()} in paragraph {i+1}:")
                        print(f"   Text: {text[:100]}...")
    
    # Analyze tables
    for table_idx, table in enumerate(doc.tables):
        print(f"\n📊 TABLE {table_idx + 1}:")
        print(f"   Dimensions: {len(table.rows)} rows × {len(table.columns)} columns")
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                if cell_text:
                    # Check for patterns in table cells
                    for pattern_name, pattern in sharedo_patterns.items():
                        matches = pattern.findall(cell_text)
                        if matches:
                            findings[pattern_name].extend(matches)
                            if row_idx == 0:  # Show header patterns
                                print(f"   Cell[{row_idx},{cell_idx}]: {cell_text[:50]}...")
    
    # Summary of findings
    print("\n" + "=" * 60)
    print("📋 SHAREDO TAG SUMMARY:")
    print("=" * 60)
    
    for pattern_name, items in findings.items():
        if items:
            unique_items = list(set(items))
            print(f"\n{pattern_name.replace('_', ' ').upper()}:")
            print(f"  Total: {len(items)} occurrences")
            print(f"  Unique: {len(unique_items)}")
            print(f"  Examples: {', '.join(unique_items[:5])}")
    
    # Show document preview
    print("\n📖 DOCUMENT PREVIEW (first 1000 chars):")
    print("-" * 40)
    full_doc_text = '\n'.join(full_text)
    print(full_doc_text[:1000])
    
    return doc, findings

if __name__ == "__main__":
    doc, findings = analyze_sharedo_doc("SUPLC1031.docx")
    print("\n✅ Analysis complete!")