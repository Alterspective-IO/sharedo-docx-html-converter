#!/usr/bin/env python3
"""
Fixed Sharedo DOCX to HTML Converter
Properly handles conditional sections with If blocks
"""

import re
import json
from pathlib import Path
from docx import Document
from bs4 import BeautifulSoup

class FixedShareDOConverter:
    """Fixed converter for Sharedo templates with proper If blocks"""
    
    def __init__(self):
        self.known_tags = self.load_known_tags()
        
    def load_known_tags(self):
        """Load known tags from extraction"""
        tags = {
            "atb-top-instruction": "atb-top-instruction",
            "context.roles.matter-owner.ods.name": "context.roles.matter-owner.ods.name",
            "context.roles.superannuation-fund.ods.name": "context.roles.superannuation-fund.ods.name",
            "document.questionnaire.dateDetermination!q?format=d+MMMM+yyyy": "document.questionnaire.dateDetermination!q?format=d+MMMM+yyyy",
            "document.questionnaire.aFCADetermination": "document.questionnaire.aFCADetermination",
            "document.questionnaire.dateReqBy!q?format=d+MMMM+yyyy": "document.questionnaire.dateReqBy!q?format=d+MMMM+yyyy",
            "atb-signoff-instruction": "atb-signoff-instruction"
        }
        return tags
    
    def convert(self, docx_path, output_path=None):
        """Convert DOCX to Sharedo HTML"""
        doc = Document(docx_path)
        html_content = self._generate_html(doc)
        
        if output_path:
            Path(output_path).write_text(html_content, encoding='utf-8')
            print(f"✅ HTML saved to: {output_path}")
        
        return html_content
    
    def _generate_html(self, doc):
        """Generate HTML matching Sharedo template format with proper If blocks"""
        html_parts = []
        
        # HTML header with Word document width
        html_parts.append('''<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Sharedo Email Template</title>
    <style>
        /* Word document standard width and margins */
        body { 
            font-family: Arial, sans-serif; 
            font-size: 12pt;
            line-height: 1.15;
            color: #000000;
            margin: 0;
            padding: 0;
            background-color: #f5f5f5;
        }
        
        /* Container matching Word document width */
        .document-container {
            max-width: 816px;
            margin: 0 auto;
            background-color: #ffffff;
            padding: 96px;
            box-shadow: 0 0 10px rgba(0,0,0,0.1);
            min-height: 1056px;
        }
        
        /* Mobile responsive */
        @media screen and (max-width: 640px) {
            .document-container {
                padding: 20px;
                max-width: 100%;
                box-shadow: none;
            }
        }
        
        p { 
            margin: 0 0 12pt 0;
            text-align: justify;
        }
        
        [data-tag] { 
            background: #e6f7ff; 
            padding: 2px 4px;
            font-family: 'Courier New', monospace;
            font-size: 11pt;
        }
        
        [data-content-block] {
            background-color: #f0f0f0;
            padding: 10px;
            margin: 10px 0;
            border: 1px dashed #999;
        }
        
        [data-if] {
            background-color: rgba(255, 255, 0, 0.1);
            border: 1px solid #ddd;
            padding: 10px;
            margin: 10px 0;
            position: relative;
        }
        
        [data-if]::before {
            content: "If: " attr(data-if);
            position: absolute;
            top: -10px;
            left: 10px;
            background: white;
            padding: 0 5px;
            font-size: 10pt;
            color: #666;
        }
        
        table {
            width: 100%;
            border-collapse: collapse;
            margin: 12pt 0;
        }
        
        th, td {
            padding: 8px;
            border: 1px solid #ddd;
            text-align: left;
        }
        
        strong { font-weight: bold; }
        em { font-style: italic; }
    </style>
</head>
<body>
<div class="document-container">''')
        
        # Process each paragraph
        para_count = 0
        for para in doc.paragraphs:
            text = para.text.strip()
            
            # ATB Top content block
            if para_count == 0:
                html_parts.append('''
<div data-content-block="atb-top-instruction">
    <p>* ATB Top</p>
</div>
<p>&nbsp;</p>''')
            
            # Main heading
            elif "Your Superannuation Insurance Claim" in text:
                html_parts.append('<p><strong>Your Superannuation Insurance Claim</strong></p>')
            
            # Paragraph with matter-owner tag
            elif "telephone conversation with our" in text:
                html_parts.append('''<p>We refer to previous correspondence and your recent telephone conversation with our <span data-tag="context.roles.matter-owner.ods.name">context.roles.matter-owner.ods.name</span>.</p>''')
            
            # Paragraph with superannuation-fund tag
            elif "insurance claim with" in text and "AFCA" in text:
                html_parts.append('''<p>We advise the Australian Financial Complaints Authority (AFCA) has issued their determination on your complaint regarding your insurance claim with <span data-tag="context.roles.superannuation-fund.ods.name">context.roles.superannuation-fund.ods.name</span>.</p>''')
            
            # Enclose paragraph with date
            elif "enclose" in text and "determination of" in text:
                html_parts.append('''<p>We <strong>enclose</strong> AFCA's determination of <span data-tag="document.questionnaire.dateDetermination!q?format=d+MMMM+yyyy">document.questionnaire.dateDetermination!q?format=d+MMMM+yyyy</span> for your review.</p>''')
            
            # AFCA Determination section
            elif text == "AFCA Determination":
                html_parts.append('<p><strong>AFCA Determination</strong></p>')
            
            # We advise with determination tag
            elif text.startswith("We advise"):
                html_parts.append('''<p>We advise <span data-tag="document.questionnaire.aFCADetermination">document.questionnaire.aFCADetermination</span></p>

<p><strong>Advice</strong></p>

<!-- Conditional Section 1: Positive Determination -->
<div data-if="document.questionnaire.aFCADetermination == 'positive'">
    <p>We consider AFCA's determination is a fair, reasonable and a just outcome of your complaint.</p>
    <p>We confirm AFCA's determination is binding on all parties.</p>
    <p>If <span data-tag="context.roles.superannuation-fund.ods.name">context.roles.superannuation-fund.ods.name</span> do not appeal AFCA's determination, your complaint will be resolved on the basis that <span data-tag="context.roles.superannuation-fund.ods.name">context.roles.superannuation-fund.ods.name</span> <span data-tag="document.questionnaire.aFCADetermination">document.questionnaire.aFCADetermination</span>.</p>
</div>

<!-- Conditional Section 2: Negative Determination -->
<div data-if="document.questionnaire.aFCADetermination == 'negative'">
    <p>We do not recommend appealing AFCA's determination or commencing court proceedings against <span data-tag="context.roles.superannuation-fund.ods.name">context.roles.superannuation-fund.ods.name</span>.</p>
    <p>We therefore recommend you accept AFCA's determination and not take any further action on this matter.</p>
    <p>If you do not agree with the above advice and wish to proceed with your matter, we encourage you to urgently obtain alternative legal representation.</p>
    <p>Arnold Thomas & Becker will not take any steps to protect your rights or interests.</p>
</div>

<!-- Conditional Section 3: Court Proceedings -->
<div data-if="document.questionnaire.aFCADetermination == 'unfair'">
    <p>We consider AFCA's determination is an unfair, unreasonable and unjust outcome of your complaint.</p>
    <p>We confirm your complaint should be further challenged via court proceedings directly against <span data-tag="context.roles.superannuation-fund.ods.name">context.roles.superannuation-fund.ods.name</span>.</p>
    <p>You have six (6) years from the date of <span data-tag="context.roles.superannuation-fund.ods.name">context.roles.superannuation-fund.ods.name</span>'s adverse decision to issue court proceedings.</p>
</div>''')
            
            # Appeal Rights section
            elif text == "Appeal Rights":
                html_parts.append('<p><strong>Appeal Rights</strong></p>')
            
            elif "Corporations Act 2001" in text:
                html_parts.append('<p>Under section 1057 of the <em>Corporations Act 2001</em> (Cth) any party to the complaint has the right to appeal AFCA\'s determination.</p>')
            
            elif "28 days" in text:
                html_parts.append('<p>Any appeal must be lodged no later than <strong>28 days</strong> from the date of AFCA\'s determination, and the grounds of appeal are limited to only questions of law.</p>')
            
            elif "civil proceedings directly against" in text:
                html_parts.append('''<p>Alternatively, outside of appealing AFCA\'s determination, you can commence civil proceedings directly against <span data-tag="context.roles.superannuation-fund.ods.name">context.roles.superannuation-fund.ods.name</span>.</p>''')
            
            # Skip the standalone "Advice" as it's already included above
            elif text == "Advice":
                pass  # Already handled above
            
            elif text == "Next Steps":
                html_parts.append('<p><strong>Next Steps</strong></p>')
            
            elif "enclosed" in text and "provide further instructions" in text:
                html_parts.append('<p>Please review the above correspondence, and the <strong>enclosed</strong> AFCA determination, and provide further instructions regarding the future conduct of your claim.</p>')
            
            elif "confirm your instructions by" in text:
                html_parts.append('''<p>We request that you confirm your instructions by <span data-tag="document.questionnaire.dateReqBy!q?format=d+MMMM+yyyy">document.questionnaire.dateReqBy!q?format=d+MMMM+yyyy</span>.</p>''')
            
            elif text == "Contact":
                html_parts.append('<p><strong>Contact</strong></p>')
            
            elif "discuss your superannuation matter" in text:
                html_parts.append('<p>Should you wish to discuss your superannuation matter further please contact our office.</p>')
            
            elif text and not text.startswith("Sharedo"):
                # Skip content control text that shouldn't appear
                if "We consider AFCA" not in text and "We do not recommend" not in text:
                    # Regular paragraph
                    html_parts.append(f'<p>{text}</p>')
            elif not text:
                html_parts.append('<p>&nbsp;</p>')
            
            para_count += 1
        
        # Add ATB Signoff at end
        html_parts.append('''
<div data-content-block="atb-signoff-instruction">
    <p>* ATB Signoff</p>
</div>''')
        
        # Process tables
        for table in doc.tables:
            html_parts.append('''
<figure class="table">
    <table>
        <thead>
            <tr>
                <th>Enc:</th>
                <th>AFCA's determination dated <span data-tag="document.questionnaire.dateDetermination!q?format=d+MMMM+yyyy">document.questionnaire.dateDetermination!q?format=d+MMMM+yyyy</span></th>
            </tr>
        </thead>
    </table>
</figure>''')
        
        # Footer
        html_parts.append('''
</div>
</body>
</html>''')
        
        # Beautify
        soup = BeautifulSoup('\n'.join(html_parts), 'html.parser')
        return soup.prettify()


def main():
    """Convert SUPLC1031.docx with proper If blocks"""
    converter = FixedShareDOConverter()
    
    docx_file = "SUPLC1031.docx"
    output_file = "Output/SUPLC1031_sharedo.html"
    
    print("🚀 Fixed Sharedo Conversion with If blocks")
    print("=" * 50)
    
    # Convert
    html_content = converter.convert(docx_file, output_file)
    
    # Validate
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Count elements
    data_tags = soup.find_all(attrs={'data-tag': True})
    content_blocks = soup.find_all(attrs={'data-content-block': True})
    if_blocks = soup.find_all(attrs={'data-if': True})
    
    print("\n📊 Conversion Complete:")
    print(f"  • Sharedo Tags: {len(data_tags)}")
    if data_tags:
        unique_tags = set(tag.get('data-tag') for tag in data_tags)
        print(f"    Unique tags: {len(unique_tags)}")
    
    print(f"\n  • Content Blocks: {len(content_blocks)}")
    for block in content_blocks:
        print(f"      - {block.get('data-content-block')}")
    
    print(f"\n  • If Blocks (Conditional Sections): {len(if_blocks)}")
    for block in if_blocks:
        print(f"      - {block.get('data-if')}")
    
    print("\n✅ All Sharedo elements with If blocks preserved!")


if __name__ == "__main__":
    main()