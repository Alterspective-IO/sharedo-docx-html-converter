#!/usr/bin/env python3
"""
Final Working Sharedo DOCX to HTML Converter
This version properly identifies and replaces content controls
"""

import re
import json
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
from docx import Document
from bs4 import BeautifulSoup

class FinalShareDOConverter:
    """Final working converter for Sharedo templates"""
    
    def __init__(self):
        self.content_control_map = {}
        self.paragraph_controls = {}
        
    def extract_control_positions(self, docx_path):
        """Extract content control positions in document"""
        with zipfile.ZipFile(docx_path, 'r') as docx_zip:
            if 'word/document.xml' in docx_zip.namelist():
                doc_xml = docx_zip.read('word/document.xml').decode('utf-8')
                
                # Register namespaces
                namespaces = {
                    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                }
                
                root = ET.fromstring(doc_xml)
                
                # Build a map of paragraph index to content controls
                paragraphs = root.findall('.//w:p', namespaces)
                
                for p_idx, para in enumerate(paragraphs):
                    # Get all text in paragraph
                    all_texts = para.findall('.//w:t', namespaces)
                    full_text = ''.join([t.text or '' for t in all_texts])
                    
                    # Find SDT elements in this paragraph
                    sdts = para.findall('.//w:sdt', namespaces)
                    
                    if sdts:
                        controls = []
                        for sdt in sdts:
                            sdt_pr = sdt.find('w:sdtPr', namespaces)
                            if sdt_pr:
                                tag_elem = sdt_pr.find('w:tag', namespaces)
                                alias_elem = sdt_pr.find('w:alias', namespaces)
                                
                                if tag_elem is not None:
                                    tag = tag_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                                    alias = alias_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') if alias_elem is not None else None
                                    
                                    controls.append({
                                        'tag': tag,
                                        'alias': alias,
                                        'type': self._get_control_type(alias)
                                    })
                        
                        if controls:
                            self.paragraph_controls[p_idx] = controls
                            
                            # Special handling for sections - they contain multiple paragraphs
                            if any(c['type'] == 'section' for c in controls):
                                # Store the full section text
                                sdt_content = para.find('.//w:sdtContent', namespaces)
                                if sdt_content:
                                    section_texts = sdt_content.findall('.//w:t', namespaces)
                                    section_text = ''.join([t.text or '' for t in section_texts])
                                    for c in controls:
                                        if c['type'] == 'section':
                                            c['content'] = section_text
        
        print(f"Found controls in {len(self.paragraph_controls)} paragraphs")
        return self.paragraph_controls
    
    def _get_control_type(self, alias):
        """Determine control type from alias"""
        if not alias:
            return 'tag'
        if 'ContentBlock' in alias:
            return 'content_block'
        if 'Section' in alias:
            return 'section'
        if 'Tag' in alias:
            return 'tag'
        return 'unknown'
    
    def convert(self, docx_path, output_path=None):
        """Convert DOCX to Sharedo HTML"""
        
        # Extract control positions
        self.extract_control_positions(docx_path)
        
        # Process document
        doc = Document(docx_path)
        html_content = self._generate_html(doc)
        
        if output_path:
            Path(output_path).write_text(html_content, encoding='utf-8')
            print(f"✅ HTML saved to: {output_path}")
        
        return html_content
    
    def _generate_html(self, doc):
        """Generate HTML with proper Sharedo elements"""
        html_parts = []
        
        # Header
        html_parts.append('''<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Sharedo Email Template</title>
    <style type="text/css">
        body { font-family: Arial, sans-serif; font-size: 14px; line-height: 1.6; color: #333; margin: 20px; }
        p { margin: 0 0 10px 0; }
        [data-tag] { background-color: #e6f7ff; padding: 2px 4px; border-radius: 2px; font-family: monospace; font-size: 0.9em; }
        [data-content-block] { background-color: #f0f0f0; padding: 10px; margin: 10px 0; border: 1px dashed #999; }
        [data-section] { border-left: 3px solid #0969da; padding-left: 10px; margin: 10px 0; }
        .table { width: 100%; border-collapse: collapse; margin: 15px 0; }
        .table th, .table td { padding: 8px; border: 1px solid #ddd; text-align: left; }
    </style>
</head>
<body>''')
        
        # Process paragraphs
        for p_idx, para in enumerate(doc.paragraphs):
            para_html = self._process_paragraph(para, p_idx)
            if para_html:
                html_parts.append(para_html)
        
        # Process tables
        for table in doc.tables:
            table_html = self._process_table(table)
            html_parts.append(table_html)
        
        # Footer
        html_parts.append('</body>\n</html>')
        
        # Beautify
        soup = BeautifulSoup('\n'.join(html_parts), 'html.parser')
        return soup.prettify()
    
    def _process_paragraph(self, para, p_idx):
        """Process paragraph with content control awareness"""
        text = para.text.strip()
        
        # Check if this paragraph has content controls
        if p_idx in self.paragraph_controls:
            controls = self.paragraph_controls[p_idx]
            
            for control in controls:
                if control['type'] == 'content_block':
                    # Content block
                    alias = control['alias'].replace('Sharedo ContentBlock: ', '') if control['alias'] else ''
                    return f'''<div data-content-block="{control['tag']}">
    <p>{alias}</p>
</div>'''
                
                elif control['type'] == 'section':
                    # Section with content
                    content = control.get('content', '')
                    # Process the section content to replace embedded tags
                    processed_content = self._process_section_content(content, control['tag'])
                    return f'''<div data-section="{control['tag']}">
{processed_content}
</div>'''
                
                elif control['type'] == 'tag':
                    # Inline tag - replace in text
                    if text:
                        # The text might have the tag embedded
                        return self._replace_inline_tags(text, para, control)
        
        # Regular paragraph
        if not text:
            return '<p>&nbsp;</p>'
        
        # Check for inline tags in regular paragraphs
        return self._format_paragraph(text, para)
    
    def _process_section_content(self, content, section_tag):
        """Process content within a section"""
        if not content:
            return ''
        
        # Split into lines and process each
        lines = content.split('\n')
        processed_lines = []
        
        for line in lines:
            line = line.strip()
            if line:
                # Look for embedded Sharedo tags
                line = self._replace_sharedo_markers(line)
                processed_lines.append(f'    <p>{line}</p>')
        
        return '\n'.join(processed_lines)
    
    def _replace_sharedo_markers(self, text):
        """Replace Sharedo tag markers with proper HTML"""
        # Pattern for Sharedo tags
        patterns = [
            (r'Sharedo Tag:\s*([a-zA-Z0-9.!?_\-=+&]+)', r'<span data-tag="\1">\1</span>'),
            (r'context\.roles\.[a-zA-Z0-9.\-_]+\.ods\.[a-zA-Z0-9.\-_]+', lambda m: f'<span data-tag="{m.group()}">{m.group()}</span>'),
            (r'context\.[a-zA-Z0-9.\-_!?=+&]+', lambda m: f'<span data-tag="{m.group()}">{m.group()}</span>'),
            (r'document\.[a-zA-Z0-9.\-_!?=+&]+', lambda m: f'<span data-tag="{m.group()}">{m.group()}</span>'),
        ]
        
        for pattern, replacement in patterns:
            text = re.sub(pattern, replacement, text)
        
        return text
    
    def _replace_inline_tags(self, text, para, control):
        """Replace inline tags in paragraph text"""
        # The paragraph contains an inline tag
        # We need to identify where it should be placed
        
        # Look at the runs to understand the structure
        formatted_parts = []
        tag_inserted = False
        
        for run in para.runs:
            run_text = run.text
            
            # If this run is empty or just a space/period after "our" or "with", insert tag
            if not tag_inserted and (
                (run_text in [' ', '.', ' .']) or 
                (not run_text.strip()) or
                ('our' in text and run_text == '.') or
                ('with' in text and run_text == '.')
            ):
                # Insert the tag
                tag_html = f'<span data-tag="{control["tag"]}">{control["tag"]}</span>'
                if run_text == '.':
                    formatted_parts.append(tag_html + '.')
                elif run_text == ' .':
                    formatted_parts.append(' ' + tag_html + '.')
                else:
                    formatted_parts.append(tag_html)
                tag_inserted = True
            else:
                # Regular text
                if run.bold:
                    run_text = f'<strong>{run_text}</strong>'
                if run.italic:
                    run_text = f'<em>{run_text}</em>'
                if run.underline:
                    run_text = f'<u>{run_text}</u>'
                formatted_parts.append(run_text)
        
        # If tag wasn't inserted yet, append it
        if not tag_inserted and control:
            tag_html = f'<span data-tag="{control["tag"]}">{control["tag"]}</span>'
            # Find the best position to insert
            result = ''.join(formatted_parts)
            if 'our .' in result:
                result = result.replace('our .', f'our {tag_html}.')
            elif 'with .' in result:
                result = result.replace('with .', f'with {tag_html}.')
            elif 'by .' in result:
                result = result.replace('by .', f'by {tag_html}.')
            elif 'of  for' in result:
                result = result.replace('of  for', f'of {tag_html} for')
            else:
                result += ' ' + tag_html
            return f'<p>{result}</p>'
        
        return f'<p>{"".join(formatted_parts)}</p>'
    
    def _format_paragraph(self, text, para):
        """Format regular paragraph"""
        # Check for formatting in runs
        if para.runs:
            formatted_parts = []
            for run in para.runs:
                run_text = run.text
                if run.bold:
                    run_text = f'<strong>{run_text}</strong>'
                if run.italic:
                    run_text = f'<em>{run_text}</em>'
                if run.underline:
                    run_text = f'<u>{run_text}</u>'
                formatted_parts.append(run_text)
            text = ''.join(formatted_parts)
        
        # Check for heading
        if para.style and para.style.name.startswith('Heading'):
            level = 2
            try:
                level = int(re.search(r'\d', para.style.name).group())
            except:
                pass
            return f'<h{level}>{text}</h{level}>'
        
        return f'<p>{text}</p>'
    
    def _process_table(self, table):
        """Process table"""
        html_parts = ['<figure class="table">\n<table>']
        
        if table.rows:
            # Header
            html_parts.append('\n    <thead>\n        <tr>')
            for cell in table.rows[0].cells:
                cell_text = cell.text.strip()
                # Check for inline tags
                cell_text = self._replace_sharedo_markers(cell_text)
                html_parts.append(f'\n            <th>{cell_text}</th>')
            html_parts.append('\n        </tr>\n    </thead>')
            
            # Body
            if len(table.rows) > 1:
                html_parts.append('\n    <tbody>')
                for row in table.rows[1:]:
                    html_parts.append('\n        <tr>')
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        cell_text = self._replace_sharedo_markers(cell_text)
                        html_parts.append(f'\n            <td>{cell_text}</td>')
                    html_parts.append('\n        </tr>')
                html_parts.append('\n    </tbody>')
        
        html_parts.append('\n</table>\n</figure>')
        return ''.join(html_parts)


def main():
    """Convert SUPLC1031.docx"""
    converter = FinalShareDOConverter()
    
    docx_file = "SUPLC1031.docx"
    output_file = "SUPLC1031_final.html"
    
    print("🚀 Final Sharedo Conversion")
    print("=" * 50)
    
    # Convert
    html_content = converter.convert(docx_file, output_file)
    
    # Validate
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Count elements
    data_tags = soup.find_all(attrs={'data-tag': True})
    content_blocks = soup.find_all(attrs={'data-content-block': True})
    sections = soup.find_all(attrs={'data-section': True})
    
    print("\n📊 Results:")
    print(f"  • Sharedo Tags: {len(data_tags)}")
    if data_tags:
        print("    Sample tags:")
        for tag in data_tags[:8]:
            print(f"      - {tag.get('data-tag')}")
    
    print(f"  • Content Blocks: {len(content_blocks)}")
    for block in content_blocks:
        print(f"      - {block.get('data-content-block')}")
    
    print(f"  • Sections: {len(sections)}")
    for section in sections:
        print(f"      - {section.get('data-section')}")
    
    print("\n✅ Conversion complete!")


if __name__ == "__main__":
    main()