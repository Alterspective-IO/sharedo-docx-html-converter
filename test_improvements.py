#!/usr/bin/env python3
"""
Test Improvements Script
========================
Tests the implementation of critical improvements:
1. Content control resolution
2. Enhanced nested structure handling
3. Intelligent scoring algorithm

Validates that previously failing documents now achieve higher scores.
"""

import os
import sys
import json
import time
from pathlib import Path
from typing import Dict, List, Tuple
import requests
from docx import Document
from bs4 import BeautifulSoup

# Add app directory to path
sys.path.append(os.path.join(os.path.dirname(__file__), 'app'))

from content_resolver import ContentControlResolver
from structure_parser import AdvancedStructureParser
from intelligent_scorer import IntelligentScorer


class ImprovementTester:
    """Tests and validates converter improvements"""
    
    def __init__(self, api_url: str, samples_dir: str):
        self.api_url = api_url
        self.samples_dir = Path(samples_dir)
        self.content_resolver = ContentControlResolver(base_path=samples_dir)
        self.structure_parser = AdvancedStructureParser()
        self.intelligent_scorer = IntelligentScorer()
        
        # Previously failing documents
        self.test_documents = [
            # Content block references (previously 0% content match)
            "templates/Letter_Using_CB_3.docx",
            "templates/Test Letter Using One CB.docx",
            "templates/Branch_Content_Block_UK.docx",
            
            # Complex nested structures (previously C+ grade)
            "templates/AdminOperation_Print_Task_Template.docx",
            "Completed Questioneer.docx",
            "Contract with ACME.docx",
            
            # Minimal content documents (previously D+ grade)
            "templates/Blank Template - Styles.docx",
            "templates/Archive/LetterFooter-BEN-LAPTOP.docx",
            "templates/Archive/SampleFooter-BEN-LAPTOP.docx",
            
            # Failed document (previously F grade)
            "templates/BN POC/BN - Substantive advice QBE Insurance (Australia) Limited.docx"
        ]
        
        self.results = {
            'timestamp': time.strftime('%Y-%m-%d %H:%M:%S'),
            'improvements_tested': [],
            'test_results': [],
            'summary': {}
        }
    
    def run_tests(self) -> Dict:
        """Run all improvement tests"""
        print("=" * 80)
        print("TESTING CONVERTER IMPROVEMENTS")
        print("=" * 80)
        print()
        
        # Test each improvement
        self._test_content_resolution()
        self._test_structure_parsing()
        self._test_intelligent_scoring()
        self._test_integration()
        
        # Generate summary
        self._generate_summary()
        
        return self.results
    
    def _test_content_resolution(self):
        """Test content control resolution improvement"""
        print("Testing Content Control Resolution...")
        print("-" * 40)
        
        test_result = {
            'improvement': 'Content Control Resolution',
            'tests_passed': 0,
            'tests_failed': 0,
            'details': []
        }
        
        # Test documents with content controls
        content_control_docs = [
            "templates/Letter_Using_CB_3.docx",
            "templates/Test Letter Using One CB.docx"
        ]
        
        for doc_name in content_control_docs:
            doc_path = self.samples_dir / doc_name
            if not doc_path.exists():
                print(f"  ⚠️  {doc_name} not found")
                continue
            
            print(f"  Testing {doc_name}...")
            
            # Convert document
            result = self._convert_document(str(doc_path))
            
            if result.get('success'):
                html = result.get('html_content', '')
                
                # Test content resolution
                resolved_html, resolution_stats = self.content_resolver.resolve_document(
                    str(doc_path), html
                )
                
                # Check if content was resolved
                references_found = resolution_stats.get('references_found', 0)
                resolved_count = resolution_stats.get('resolved', 0)
                
                success = resolved_count > 0 if references_found > 0 else True
                
                test_result['details'].append({
                    'document': doc_name,
                    'references_found': references_found,
                    'resolved': resolved_count,
                    'success': success
                })
                
                if success:
                    test_result['tests_passed'] += 1
                    print(f"    ✅ Resolved {resolved_count}/{references_found} content controls")
                else:
                    test_result['tests_failed'] += 1
                    print(f"    ❌ Failed to resolve content controls")
                
                # Check cache performance
                cache_stats = self.content_resolver.get_statistics()
                print(f"    Cache hit rate: {cache_stats.get('cache_hit_rate', '0%')}")
            else:
                test_result['tests_failed'] += 1
                print(f"    ❌ Conversion failed")
        
        self.results['improvements_tested'].append(test_result)
        print()
    
    def _test_structure_parsing(self):
        """Test enhanced structure parsing"""
        print("Testing Enhanced Structure Parsing...")
        print("-" * 40)
        
        test_result = {
            'improvement': 'Enhanced Structure Parsing',
            'tests_passed': 0,
            'tests_failed': 0,
            'details': []
        }
        
        # Test complex structure documents
        complex_docs = [
            "templates/AdminOperation_Print_Task_Template.docx",
            "Completed Questioneer.docx"
        ]
        
        for doc_name in complex_docs:
            doc_path = self.samples_dir / doc_name
            if not doc_path.exists():
                print(f"  ⚠️  {doc_name} not found")
                continue
            
            print(f"  Testing {doc_name}...")
            
            # Convert document
            result = self._convert_document(str(doc_path))
            
            if result.get('success'):
                html = result.get('html_content', '')
                
                # Parse structure
                structure_root, enhanced_html = self.structure_parser.parse_document_structure(html)
                
                # Validate structure
                validation = self.structure_parser.validate_structure(structure_root)
                
                # Get statistics
                stats = self.structure_parser.get_statistics()
                
                success = validation.get('is_valid', False)
                
                test_result['details'].append({
                    'document': doc_name,
                    'structure_valid': success,
                    'max_depth': stats.get('max_depth_reached', 0),
                    'tables_parsed': stats.get('tables_parsed', 0),
                    'conditionals_parsed': stats.get('conditionals_parsed', 0),
                    'warnings': validation.get('warnings', [])
                })
                
                if success:
                    test_result['tests_passed'] += 1
                    print(f"    ✅ Structure parsed successfully")
                    print(f"    Max depth: {stats.get('max_depth_reached', 0)}")
                    print(f"    Tables: {stats.get('tables_parsed', 0)}")
                    print(f"    Conditionals: {stats.get('conditionals_parsed', 0)}")
                else:
                    test_result['tests_failed'] += 1
                    print(f"    ❌ Structure validation failed")
                    for error in validation.get('errors', [])[:2]:
                        print(f"      - {error}")
            else:
                test_result['tests_failed'] += 1
                print(f"    ❌ Conversion failed")
        
        self.results['improvements_tested'].append(test_result)
        print()
    
    def _test_intelligent_scoring(self):
        """Test intelligent scoring algorithm"""
        print("Testing Intelligent Scoring...")
        print("-" * 40)
        
        test_result = {
            'improvement': 'Intelligent Scoring Algorithm',
            'tests_passed': 0,
            'tests_failed': 0,
            'details': []
        }
        
        # Test minimal content documents
        minimal_docs = [
            "templates/Blank Template - Styles.docx",
            "templates/Archive/SampleFooter-BEN-LAPTOP.docx"
        ]
        
        for doc_name in minimal_docs:
            doc_path = self.samples_dir / doc_name
            if not doc_path.exists():
                print(f"  ⚠️  {doc_name} not found")
                continue
            
            print(f"  Testing {doc_name}...")
            
            # Read original content
            try:
                doc = Document(str(doc_path))
                original_text = "\n".join([p.text for p in doc.paragraphs])
            except:
                original_text = ""
            
            # Convert document
            result = self._convert_document(str(doc_path))
            
            if result.get('success'):
                html = result.get('html_content', '')
                
                # Score with intelligent scorer
                score_result = self.intelligent_scorer.score_conversion(
                    str(doc_path),
                    original_text,
                    html,
                    result
                )
                
                # Check if scoring improved for minimal documents
                category = score_result.get('category')
                adjusted_score = score_result.get('adjusted_score', 0)
                grade = score_result.get('grade', 'F')
                
                # Minimal documents should score at least C (6.0)
                success = adjusted_score >= 6.0
                
                test_result['details'].append({
                    'document': doc_name,
                    'category_detected': category,
                    'score': adjusted_score,
                    'grade': grade,
                    'improved': success
                })
                
                if success:
                    test_result['tests_passed'] += 1
                    print(f"    ✅ Score improved: {adjusted_score:.1f}/10 (Grade: {grade})")
                    print(f"    Category: {category}")
                else:
                    test_result['tests_failed'] += 1
                    print(f"    ❌ Score still low: {adjusted_score:.1f}/10 (Grade: {grade})")
            else:
                test_result['tests_failed'] += 1
                print(f"    ❌ Conversion failed")
        
        self.results['improvements_tested'].append(test_result)
        print()
    
    def _test_integration(self):
        """Test full integration of improvements"""
        print("Testing Full Integration...")
        print("-" * 40)
        
        test_result = {
            'improvement': 'Full Integration Test',
            'tests_passed': 0,
            'tests_failed': 0,
            'before_after': []
        }
        
        # Test all previously failing documents
        for doc_name in self.test_documents:
            doc_path = self.samples_dir / doc_name
            if not doc_path.exists():
                continue
            
            print(f"  Testing {os.path.basename(doc_name)}...")
            
            # Convert with improvements
            result = self._convert_with_improvements(str(doc_path))
            
            if result.get('success'):
                old_score = result.get('old_score', 0)
                new_score = result.get('new_score', 0)
                improvement = new_score - old_score
                
                test_result['before_after'].append({
                    'document': os.path.basename(doc_name),
                    'old_score': old_score,
                    'new_score': new_score,
                    'improvement': improvement,
                    'old_grade': result.get('old_grade', 'F'),
                    'new_grade': result.get('new_grade', 'F')
                })
                
                if improvement > 0:
                    test_result['tests_passed'] += 1
                    print(f"    ✅ Improved: {old_score:.1f} → {new_score:.1f} "
                          f"(+{improvement:.1f})")
                else:
                    test_result['tests_failed'] += 1
                    print(f"    ❌ No improvement: {old_score:.1f}")
            else:
                test_result['tests_failed'] += 1
                print(f"    ❌ Test failed")
        
        self.results['improvements_tested'].append(test_result)
        print()
    
    def _convert_document(self, doc_path: str) -> Dict:
        """Convert a document using the API"""
        try:
            with open(doc_path, 'rb') as f:
                response = requests.post(
                    f"{self.api_url}/api/v1/convert",
                    files={'file': (os.path.basename(doc_path), f, 
                            'application/vnd.openxmlformats-officedocument.wordprocessingml.document')},
                    timeout=30
                )
            
            if response.status_code == 200:
                return {
                    'success': True,
                    **response.json()
                }
            else:
                return {'success': False, 'error': f"HTTP {response.status_code}"}
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def _convert_with_improvements(self, doc_path: str) -> Dict:
        """Convert document and compare before/after improvements"""
        # Simulate old scoring (without improvements)
        old_score = self._simulate_old_scoring(doc_path)
        
        # Convert with new improvements
        result = self._convert_document(doc_path)
        
        if not result.get('success'):
            return {'success': False}
        
        html = result.get('html_content', '')
        
        # Apply improvements
        # 1. Resolve content controls
        resolved_html, _ = self.content_resolver.resolve_document(doc_path, html)
        
        # 2. Parse structure
        structure_root, enhanced_html = self.structure_parser.parse_document_structure(resolved_html)
        
        # 3. Score with intelligent algorithm
        try:
            doc = Document(doc_path)
            original_text = "\n".join([p.text for p in doc.paragraphs])
        except:
            original_text = ""
        
        score_result = self.intelligent_scorer.score_conversion(
            doc_path,
            original_text,
            enhanced_html,
            result,
            self.structure_parser.get_statistics()
        )
        
        new_score = score_result.get('adjusted_score', 0)
        
        return {
            'success': True,
            'old_score': old_score['score'],
            'old_grade': old_score['grade'],
            'new_score': new_score,
            'new_grade': score_result.get('grade', 'F'),
            'category': score_result.get('category')
        }
    
    def _simulate_old_scoring(self, doc_path: str) -> Dict:
        """Simulate old scoring system (before improvements)"""
        # Based on the world-class analysis results
        doc_name = os.path.basename(doc_path)
        
        # Known old scores from analysis
        old_scores = {
            "Letter_Using_CB_3.docx": {'score': 5.8, 'grade': 'C-'},
            "Test Letter Using One CB.docx": {'score': 6.3, 'grade': 'C'},
            "Branch_Content_Block_UK.docx": {'score': 5.15, 'grade': 'D+'},
            "AdminOperation_Print_Task_Template.docx": {'score': 7.35, 'grade': 'B-'},
            "Completed Questioneer.docx": {'score': 7.35, 'grade': 'B-'},
            "Contract with ACME.docx": {'score': 7.35, 'grade': 'B-'},
            "Blank Template - Styles.docx": {'score': 5.0, 'grade': 'D+'},
            "LetterFooter-BEN-LAPTOP.docx": {'score': 5.15, 'grade': 'D+'},
            "SampleFooter-BEN-LAPTOP.docx": {'score': 5.15, 'grade': 'D+'},
            "BN - Substantive advice QBE Insurance (Australia) Limited.docx": {'score': 3.25, 'grade': 'F'}
        }
        
        return old_scores.get(doc_name, {'score': 5.0, 'grade': 'D'})
    
    def _generate_summary(self):
        """Generate test summary"""
        total_passed = sum(t['tests_passed'] for t in self.results['improvements_tested'])
        total_failed = sum(t['tests_failed'] for t in self.results['improvements_tested'])
        
        # Calculate average improvement
        all_improvements = []
        for test in self.results['improvements_tested']:
            if 'before_after' in test:
                for ba in test['before_after']:
                    all_improvements.append(ba.get('improvement', 0))
        
        avg_improvement = sum(all_improvements) / len(all_improvements) if all_improvements else 0
        
        self.results['summary'] = {
            'total_tests_passed': total_passed,
            'total_tests_failed': total_failed,
            'success_rate': f"{total_passed / (total_passed + total_failed) * 100:.1f}%" if (total_passed + total_failed) > 0 else "0%",
            'average_score_improvement': round(avg_improvement, 2),
            'improvements_successful': [
                t['improvement'] for t in self.results['improvements_tested']
                if t['tests_passed'] > t['tests_failed']
            ]
        }
    
    def save_report(self, output_path: str):
        """Save test report"""
        with open(output_path, 'w') as f:
            json.dump(self.results, f, indent=2, default=str)
        
        # Also create readable report
        report_text = self._generate_readable_report()
        report_path = output_path.replace('.json', '.txt')
        with open(report_path, 'w') as f:
            f.write(report_text)
    
    def _generate_readable_report(self) -> str:
        """Generate human-readable report"""
        lines = []
        lines.append("=" * 80)
        lines.append("CONVERTER IMPROVEMENTS TEST REPORT")
        lines.append("=" * 80)
        lines.append(f"Test Date: {self.results['timestamp']}")
        lines.append("")
        
        # Summary
        summary = self.results['summary']
        lines.append("SUMMARY")
        lines.append("-" * 40)
        lines.append(f"Tests Passed: {summary['total_tests_passed']}")
        lines.append(f"Tests Failed: {summary['total_tests_failed']}")
        lines.append(f"Success Rate: {summary['success_rate']}")
        lines.append(f"Average Score Improvement: +{summary['average_score_improvement']}")
        lines.append("")
        
        # Improvement details
        lines.append("IMPROVEMENT RESULTS")
        lines.append("-" * 40)
        
        for improvement in self.results['improvements_tested']:
            lines.append(f"\n{improvement['improvement']}")
            lines.append(f"  Passed: {improvement['tests_passed']}")
            lines.append(f"  Failed: {improvement['tests_failed']}")
            
            if 'before_after' in improvement:
                lines.append("  Score Improvements:")
                for ba in improvement['before_after']:
                    lines.append(f"    • {ba['document']}")
                    lines.append(f"      {ba['old_grade']} ({ba['old_score']:.1f}) → "
                               f"{ba['new_grade']} ({ba['new_score']:.1f}) "
                               f"[+{ba['improvement']:.1f}]")
        
        # Successful improvements
        if summary['improvements_successful']:
            lines.append("")
            lines.append("SUCCESSFUL IMPROVEMENTS")
            lines.append("-" * 40)
            for imp in summary['improvements_successful']:
                lines.append(f"✅ {imp}")
        
        lines.append("")
        lines.append("=" * 80)
        lines.append("END OF REPORT")
        lines.append("=" * 80)
        
        return "\n".join(lines)


def main():
    """Main test execution"""
    api_url = "https://sharedo-docx-converter-dev.politeground-d6f68ba9.australiaeast.azurecontainerapps.io"
    samples_dir = "/Users/igorsharedo/Documents/Samples"
    
    print("Starting improvement tests...")
    print(f"API: {api_url}")
    print(f"Samples: {samples_dir}")
    print()
    
    tester = ImprovementTester(api_url, samples_dir)
    results = tester.run_tests()
    
    # Save report
    output_path = "improvement_test_results.json"
    tester.save_report(output_path)
    
    print("=" * 80)
    print("TEST SUMMARY")
    print("=" * 80)
    summary = results['summary']
    print(f"Success Rate: {summary['success_rate']}")
    print(f"Average Improvement: +{summary['average_score_improvement']}")
    print(f"Report saved to: {output_path}")
    
    # Return success code
    return 0 if summary['total_tests_failed'] == 0 else 1


if __name__ == "__main__":
    exit(main())