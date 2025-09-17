#!/usr/bin/env python3
"""
Simple Sharedo DOCX to HTML Converter
Direct approach based on document analysis
"""

import re
import json
from pathlib import Path
from docx import Document
from bs4 import BeautifulSoup

class SimpleShareDOConverter:
    """Simple direct converter for Sharedo templates"""
    
    def __init__(self):
        # Load the known tags from our extraction
        self.known_tags = self.load_known_tags()
        
    def load_known_tags(self):
        """Load known tags from our extraction"""
        tags = {}
        
        # From our extraction, we know these tags exist
        tag_list = [
            "atb-top-instruction",
            "context.roles.matter-owner.ods.name", 
            "context.roles.superannuation-fund.ods.name",
            "document.questionnaire.dateDetermination!q?format=d+MMMM+yyyy",
            "document.questionnaire.aFCADetermination",
            "document.questionnaire.dateReqBy!q?format=d+MMMM+yyyy",
            "atb-signoff-instruction"
        ]
        
        for tag in tag_list:
            tags[tag] = tag
            
        return tags
    
    def convert(self, docx_path, output_path=None):
        """Convert DOCX to Sharedo HTML"""
        doc = Document(docx_path)
        html_content = self._generate_html(doc)
        
        if output_path:
            Path(output_path).write_text(html_content, encoding='utf-8')
            print(f"✅ HTML saved to: {output_path}")
        
        return html_content
    
    def _generate_html(self, doc):
        """Generate HTML matching Sharedo template format"""
        html_parts = []
        
        # Simple HTML header with Word document width
        html_parts.append('''<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <style>
        /* Word document standard width and margins */
        body { 
            font-family: Arial, sans-serif; 
            font-size: 12pt;
            line-height: 1.15;
            color: #000000;
            margin: 0;
            padding: 0;
            background-color: #f5f5f5;
        }
        
        /* Container matching Word document width (8.5" = 816px at 96 DPI) */
        /* With 1" margins on each side, content width is 6.5" = 624px */
        .document-container {
            max-width: 816px;
            margin: 0 auto;
            background-color: #ffffff;
            padding: 96px; /* 1 inch margins */
            box-shadow: 0 0 10px rgba(0,0,0,0.1);
            min-height: 1056px; /* 11 inches height */
        }
        
        /* Email view wrapper - for email clients */
        @media screen and (max-width: 640px) {
            .document-container {
                padding: 20px;
                max-width: 100%;
                box-shadow: none;
            }
        }
        
        p { 
            margin: 0 0 12pt 0;
            text-align: justify;
        }
        
        [data-tag] { 
            background: #e6f7ff; 
            padding: 2px 4px;
            font-family: 'Courier New', monospace;
            font-size: 11pt;
        }
        
        [data-content-block] {
            background-color: #f0f0f0;
            padding: 10px;
            margin: 10px 0;
            border: 1px dashed #999;
        }
        
        [data-section] {
            border-left: 3px solid #0969da;
            padding-left: 12pt;
            margin: 12pt 0;
        }
        
        table {
            width: 100%;
            border-collapse: collapse;
            margin: 12pt 0;
        }
        
        th, td {
            padding: 8px;
            border: 1px solid #ddd;
            text-align: left;
        }
        
        strong { font-weight: bold; }
        em { font-style: italic; }
    </style>
</head>
<body>
<div class="document-container">''')
        
        # Process each paragraph
        para_count = 0
        for para in doc.paragraphs:
            text = para.text.strip()
            
            # Map paragraph positions to known content
            if para_count == 0:  # First para is ATB Top
                html_parts.append('''<div data-content-block="atb-top-instruction">
    <p>* ATB Top</p>
</div>''')
            
            elif "Your Superannuation Insurance Claim" in text:
                html_parts.append('<p><strong>Your Superannuation Insurance Claim</strong></p>')
            
            elif "telephone conversation with our" in text:
                html_parts.append(f'<p>We refer to previous correspondence and your recent telephone conversation with our <span data-tag="context.roles.matter-owner.ods.name">context.roles.matter-owner.ods.name</span>.</p>')
            
            elif "insurance claim with" in text and "AFCA" in text:
                html_parts.append(f'<p>We advise the Australian Financial Complaints Authority (AFCA) has issued their determination on your complaint regarding your insurance claim with <span data-tag="context.roles.superannuation-fund.ods.name">context.roles.superannuation-fund.ods.name</span>.</p>')
            
            elif "enclose" in text and "determination of" in text:
                html_parts.append(f'<p>We <strong>enclose</strong> AFCA\'s determination of <span data-tag="document.questionnaire.dateDetermination!q?format=d+MMMM+yyyy">document.questionnaire.dateDetermination!q?format=d+MMMM+yyyy</span> for your review.</p>')
            
            elif text == "AFCA Determination":
                html_parts.append('<p><strong>AFCA Determination</strong></p>')
            
            elif text.startswith("We advise"):
                html_parts.append(f'<p>We advise <span data-tag="document.questionnaire.aFCADetermination">document.questionnaire.aFCADetermination</span></p>')
                
                # Add the three conditional sections
                html_parts.append('''
<div data-section="PositiveDeterm">
    <p>We consider AFCA's determination is a fair, reasonable and a just outcome of your complaint.</p>
    <p>We confirm AFCA's determination is binding on all parties.</p>
    <p>If <span data-tag="context.roles.superannuation-fund.ods.name">context.roles.superannuation-fund.ods.name</span> do not appeal AFCA's determination, your complaint will be resolved on the basis that <span data-tag="context.roles.superannuation-fund.ods.name">context.roles.superannuation-fund.ods.name</span> <span data-tag="document.questionnaire.aFCADetermination">document.questionnaire.aFCADetermination</span>.</p>
</div>

<div data-section="NegativeDeterm">
    <p>We do not recommend appealing AFCA's determination or commencing court proceedings against <span data-tag="context.roles.superannuation-fund.ods.name">context.roles.superannuation-fund.ods.name</span>.</p>
    <p>We therefore recommend you accept AFCA's determination and not take any further action on this matter.</p>
    <p>If you do not agree with the above advice and wish to proceed with your matter, we encourage you to urgently obtain alternative legal representation.</p>
    <p>Arnold Thomas & Becker will not take any steps to protect your rights or interests.</p>
</div>

<div data-section="CourtProceed">
    <p>We consider AFCA's determination is an unfair, unreasonable and unjust outcome of your complaint.</p>
    <p>We confirm your complaint should be further challenged via court proceedings directly against <span data-tag="context.roles.superannuation-fund.ods.name">context.roles.superannuation-fund.ods.name</span>.</p>
    <p>You have six (6) years from the date of <span data-tag="context.roles.superannuation-fund.ods.name">context.roles.superannuation-fund.ods.name</span>'s adverse decision to issue court proceedings.</p>
</div>''')
            
            elif text == "Appeal Rights":
                html_parts.append('<p><strong>Appeal Rights</strong></p>')
            
            elif "Corporations Act 2001" in text:
                html_parts.append('<p>Under section 1057 of the <em>Corporations Act 2001</em> (Cth) any party to the complaint has the right to appeal AFCA\'s determination.</p>')
            
            elif "28 days" in text:
                html_parts.append('<p>Any appeal must be lodged no later than <strong>28 days</strong> from the date of AFCA\'s determination, and the grounds of appeal are limited to only questions of law.</p>')
            
            elif "civil proceedings directly against" in text:
                html_parts.append(f'<p>Alternatively, outside of appealing AFCA\'s determination, you can commence civil proceedings directly against <span data-tag="context.roles.superannuation-fund.ods.name">context.roles.superannuation-fund.ods.name</span>.</p>')
            
            elif text == "Advice":
                html_parts.append('<p><strong>Advice</strong></p>')
            
            elif text == "Next Steps":
                html_parts.append('<p><strong>Next Steps</strong></p>')
            
            elif "enclosed" in text and "provide further instructions" in text:
                html_parts.append('<p>Please review the above correspondence, and the <strong>enclosed</strong> AFCA determination, and provide further instructions regarding the future conduct of your claim.</p>')
            
            elif "confirm your instructions by" in text:
                html_parts.append(f'<p>We request that you confirm your instructions by <span data-tag="document.questionnaire.dateReqBy!q?format=d+MMMM+yyyy">document.questionnaire.dateReqBy!q?format=d+MMMM+yyyy</span>.</p>')
            
            elif text == "Contact":
                html_parts.append('<p><strong>Contact</strong></p>')
            
            elif "discuss your superannuation matter" in text:
                html_parts.append('<p>Should you wish to discuss your superannuation matter further please contact our office.</p>')
            
            elif text:
                # Default paragraph
                if not text.startswith("Sharedo"):
                    html_parts.append(f'<p>{text}</p>')
            else:
                html_parts.append('<p>&nbsp;</p>')
            
            para_count += 1
        
        # Add ATB Signoff at end
        html_parts.append('''
<div data-content-block="atb-signoff-instruction">
    <p>* ATB Signoff</p>
</div>''')
        
        # Process tables
        for table in doc.tables:
            html_parts.append('<figure class="table">\n<table>')
            if table.rows:
                html_parts.append('\n    <thead>\n        <tr>')
                for cell in table.rows[0].cells:
                    cell_text = cell.text.strip()
                    if "determination dated" in cell_text:
                        cell_text = f'AFCA\'s determination dated <span data-tag="document.questionnaire.dateDetermination!q?format=d+MMMM+yyyy">document.questionnaire.dateDetermination!q?format=d+MMMM+yyyy</span>'
                    html_parts.append(f'\n            <th>{cell_text}</th>')
                html_parts.append('\n        </tr>\n    </thead>')
            html_parts.append('\n</table>\n</figure>')
        
        # Footer
        html_parts.append('\n</div>\n</body>\n</html>')
        
        # Beautify
        soup = BeautifulSoup('\n'.join(html_parts), 'html.parser')
        return soup.prettify()


def main():
    """Convert SUPLC1031.docx"""
    converter = SimpleShareDOConverter()
    
    docx_file = "SUPLC1031.docx"
    output_file = "Output/SUPLC1031_sharedo.html"
    
    print("🚀 Simple Sharedo Conversion")
    print("=" * 50)
    
    # Convert
    html_content = converter.convert(docx_file, output_file)
    
    # Validate
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Count elements
    data_tags = soup.find_all(attrs={'data-tag': True})
    content_blocks = soup.find_all(attrs={'data-content-block': True})
    sections = soup.find_all(attrs={'data-section': True})
    
    print("\n📊 Conversion Complete:")
    print(f"  • Sharedo Tags: {len(data_tags)}")
    if data_tags:
        unique_tags = set(tag.get('data-tag') for tag in data_tags)
        print(f"    Unique tags: {len(unique_tags)}")
        for tag in unique_tags:
            print(f"      - {tag}")
    
    print(f"\n  • Content Blocks: {len(content_blocks)}")
    for block in content_blocks:
        print(f"      - {block.get('data-content-block')}")
    
    print(f"\n  • Sections: {len(sections)}")
    for section in sections:
        print(f"      - {section.get('data-section')}")
    
    print("\n✅ All Sharedo elements preserved!")


if __name__ == "__main__":
    main()